
import shutil
import copy
import mome
import subprocess
import inspect
import re
import os
import json
import glob
from docx import Document
import momeutils


def enhance_report_section(section_identifier = None, current_contents = None, initial_global_contents = None, subsection_titles = None): 
    
    icl_examples = momeutils.load_icl(inspect.currentframe())
    
    
    messages = [
            {"role": "system", "content": """
    You are a content enhancement specialist with expertise in report writing. Your task is to enhance the {target} subsection of a report using provided the initial overview of the contents of the parent section and the existing subsection content..
    You will receive the global contents, the list of titles for the subsections within the section (to help you trim and focus on the key part) and the current contents of the subsection. Your goal is to rewrite and enhance the content for the {target} subsection.
    Answer in a valid JSON format as follows:
    ```json
    {{
    
        "enhanced_section_contents" : "The enhanced contents of the {target} subsection after processing."
    
    }} 
    ```
    """.format(target = section_identifier)}, 
        {"role": "user", "content": " {} For this particular task instance, the following elements are provided:\n Current Global Contents\n{}\n\n Existing Subsections: {}\n\n Initial Subsection Contents\n{}\n\nEnhance the {} subsection using the provided information.\n\n".format(icl_examples, initial_global_contents, subsection_titles, current_contents, section_identifier)}
    ]
    
    # results = momeutils.parse_json(momeutils.ask_llm(messages, model = "g4o"))
    parseable = False
    while not parseable: 
        initial_answer, parseable, results = momeutils.safe_llm_ask(messages, model = 'g4o') 
    momeutils.crprint(json.dumps(results, indent = 4))
    return results["enhanced_section_contents"]


def rate_key_concepts_relatedness(initial_content = None, subsections = None, target_subsection = None, current_extracted_text = None,key_concepts = None): 
    
    icl_examples = momeutils.load_icl(inspect.currentframe())
    
    
    messages = [
            {"role": "system", "content": """
    You are a content analysis expert, specializing in evaluating the relatedness of key concepts to specific subsections of text. Your task is to analyze a provided text, identify its subsections, and then rate the relatedness of each extracted key concept to the ** {subs} ** subsection. You will receive the following inputs:
    
    * **Initial Text:** The full text to be analyzed.
    * **Subsections:** A list of all subsections identified within the text to enhance contex awareness and ground rating.
    * **Target Subsection Text:** The text content of the target subsection.
    * **Key Concepts:** A list of key concepts extracted from the text.
             
    Answer in a valid JSON format as follows:
    ```json
    {{
    
      
        {scores}
        
    
    }} 
    ```
    """.format(subs = target_subsection,
               scores = ",\n".join(['"{}" : int score'.format(kc.lower().replace(' ', '_')) for kc in key_concepts]))},
        {"role": "user", "content": " {} For this particular task instance, the following elements are provided:\n Initial Content\n{}\n\n Subsections\n{}\n\n Current Extracted Text\n{} Key Concepts\n{}\n\n\n\nEvaluate and rate the relatedness of the provided key concepts to the ** {} ** subsection.\n\n".format(icl_examples, initial_content, subsections, current_extracted_text, key_concepts, target_subsection)}
    ]
    
    # results = momeutils.parse_json(momeutils.ask_llm(messages, model = "g4o"))
    parseable = False
    while not parseable: 
        initial_answer, parseable, results = momeutils.safe_llm_ask(messages, model = 'g4o') 
    momeutils.crprint(json.dumps(results, indent = 4))
    return results


def extract_key_concepts(text = None): 
    
    icl_examples = momeutils.load_icl(inspect.currentframe())
    
    
    messages = [
            {"role": "system", "content": """
    You are a content summarizer with expertise in identifying key concepts. Your task is to extract the most important ideas from a given text and present them as a list. Focus on identifying the core themes and central arguments, avoiding unnecessary details or peripheral information.
    Answer in a valid JSON format as follows:
    ```json
    {{
    
        "key_concepts" : [
            "key concept 0 " ,
            "key concept 1" ,
            // add more items as needed...
        ] 
     
    }} 
    ```
    """}, 
        {"role": "user", "content": " {} For this particular task instance, the following elements are provided:\n Text\n{}\n\nExtract key concepts or ideas from the provided text and return them in a list under the key 'key_concepts'.\n\n".format(icl_examples, text)}
    ]
    
    # results = momeutils.parse_json(momeutils.ask_llm(messages, model = "g4o"))
    parseable = False
    while not parseable: 
        initial_answer, parseable, results = momeutils.safe_llm_ask(messages, model = 'g4o') 
    momeutils.crprint(json.dumps(results, indent = 4))
    return results["key_concepts"]

def extract_complementary_info_concepts(text = None, key_concepts = None): 
    
    icl_examples = momeutils.load_icl(inspect.currentframe())
    
    
    messages = [
            {"role": "system", "content": """
    You are a content summarizer with strong and robust expertise. Your task is to extract complementary information from a given text and present them as a list. 
    Specifically, you'll be provided with some initial content, the key concepts previously extracted and you must focus on identifying names, context, acronyms, methods, strategy or paradigms (only when applicable, if there are none, do not add) that are to be preserved and passed through to next sections so downstream consistency can be ensured. 
    Avoid trivial or redundant information, focusing on elements that are crucial for maintaining coherence and understanding.
    Answer in a valid JSON format as follows:
    ```json
    {{
    
        "complementary_information": [
             // names, strategy, acronyms, methods, paradigms... that might be lost downstream, make sure each item contains enough rationale to be understood in later stages
             ]
    
    }} 
    ```
    """}, 
        {"role": "user", "content": " {} For this particular task instance, the following elements are provided:\n Text\n{}\n\nExtracted key concepts\n {}\n\n Extract the remaining complementary information from the provided text that will be central in ensuring consistency and return them in a list under the key 'complementary_information'.\n\n".format(icl_examples, text, key_concepts)}
    ]
    
    # results = momeutils.parse_json(momeutils.ask_llm(messages, model = "g4o"))
    parseable = False
    while not parseable: 
        initial_answer, parseable, results = momeutils.safe_llm_ask(messages, model = 'g4o') 
    momeutils.crprint(json.dumps(results, indent = 4))
    return results["complementary_information"]



def adjust_text_based_on_review(initial_text = None, reviews = None): 
    
    icl_examples = momeutils.load_icl(inspect.currentframe())
    
    
    messages = [
            {"role": "system", "content": """
    You are a text editor with expertise in refining written content based on user feedback. Your task is to modify the initial text based on user feedback. You will receive the original text and user feedback, and your output will be the adjusted text incorporating the feedback.
    Answer in a valid JSON format as follows:
    ```json
    {{
    
        "reviewed" : "The adjusted text based on the user input and initial piece of text."
    
    }} 
    ```
    """}, 
        {"role": "user", "content": " {} For this particular task instance, the following elements are provided:\n Initial Text\n{}\n\n User reviews\n{}\n\nAdjust the initial text according to the user review provided.\n\n".format(icl_examples, initial_text, reviews)}
    ]
    
    # results = momeutils.parse_json(momeutils.ask_llm(messages, model = "g4o"))
    parseable = False
    while not parseable: 
        initial_answer, parseable, results = momeutils.safe_llm_ask(messages, model = 'g4o') 
    momeutils.crprint(json.dumps(results, indent = 4))
    return results["reviewed"]

def evaluate_parameter_relevance(text_content = None, user_example = None): 
    
    icl_examples = momeutils.load_icl(inspect.currentframe())
    
    
    messages = [
            {"role": "system", "content": """
    You are a content relevance assessor, tasked with determining the relevance of a user-provided content in a specific key to the high-level content of a text. You will receive a user-provided example and are expected to return an integer representing the relevance of the example to the text's content. Higher numbers indicate greater relevance.
    Answer in a valid JSON format as follows:
    ```json
    {{
    
        "relevance" : "An integer representing the relevance of the user-provided content regarding the section contents."
    
    }} 
    ```
    """}, 
        {"role": "user", "content": " {} For this particular task instance, the following elements are provided:\n Text Content\n{}\n\n User Example\n{}\n\nAssess the relevance of the user-provided content to the section contents and return a relevance score between 0 and 10.\n\n".format(icl_examples, text_content, user_example)}
    ]
    
    # results = momeutils.parse_json(momeutils.ask_llm(messages, model = "g4o"))
    parseable = False
    while not parseable: 
        initial_answer, parseable, results = momeutils.safe_llm_ask(messages, model = 'g4o') 
    momeutils.crprint(json.dumps(results, indent = 4))
    return int(results["relevance"])


def enhance_section_content(high_level_text = None, section_names = None, target_section = None, initial_section_content = None, user_instruction = None): 
    
    icl_examples = momeutils.load_icl(inspect.currentframe())
    
    
    messages = [
            {"role": "system", "content": """
             
    You are a content enhancement specialist with expertise in rewriting and improving text. Your task is to enhance the content of a specific section based on user instructions. You will receive the initial content of the target section and instructions for enhancing it. Your goal is to rewrite and enhance the content, making it more engaging, informative, and relevant to the user's needs.
    Answer in a valid JSON format as follows:
    ```json
    {{
    
        "enhanced_text" : "The rewritten and enhanced version of the content for the target section."
    
    }} 
    ```
    """}, 
        {"role": "user", "content": " {} For this particular task instance, the following elements are provided:\n High Level Text\n{}\n\n Section Names\n{}\n\n Target Section\n{}\n\n Initial Section Content\n{}\n\n User Instruction\n{}\n\nRewrite and enhance the content for the target section using the provided instructions.\n\n".format(icl_examples, high_level_text, section_names, target_section, initial_section_content, user_instruction)}
    ]
    
    # results = momeutils.parse_json(momeutils.ask_llm(messages, model = "g4o"))
    parseable = False
    while not parseable: 
        initial_answer, parseable, results = momeutils.safe_llm_ask(messages, model = 'g4o') 
    momeutils.crprint(json.dumps(results, indent = 4))
    return results["enhanced_text"]


def assess_text_sufficiency(core_text = None, target_subsections = None): 
    
    icl_examples = momeutils.load_icl(inspect.currentframe())
    
    
    messages = [
            {"role": "system", "content": """
    You are an expert content analyst assisting the user in evaluating the completeness of a text against a set of target subsections. Your role is to receive the text and the list of subsections and evaluate how relevant (score ranging from 0 to 10, 10 being the best match) the core text is to the distinct subsections.  
    Answer in a valid JSON format as follows:
    ```json
    {{
    
        "result": [

{subsections}
]
    
    }} 
    ```
    """.format(subsections = ",\n".join(['"Integer score for subsection {}"'.format(s) for s in target_subsections]) )}, 
        {"role": "user", "content": " {} For this particular task instance, the following elements are provided:\n Core Text\n{}\n\n Target Subsections\n{}\n\nGiven a core text and a list of target subsections, rate the relevance between the core text, returning integer scores that reflect if it contains enough information to effectively discuss each subsection.\n\n".format(icl_examples, core_text, "\n*".join([""] + target_subsections))}
    ]


    # results = momeutils.parse_json(momeutils.ask_llm(messages, model = "g4o"))
    parseable = False
    while not parseable: 
        initial_answer, parseable, results = momeutils.safe_llm_ask(messages, model = 'g4o') 
    momeutils.crprint(json.dumps(results, indent = 4))
    return results["result"]


def identify_missing_elements_for_consistency(initial_content = None, subsections = None, target_subsection = None, current_extracted_text = None): 
    
    icl_examples = momeutils.load_icl(inspect.currentframe())
    
    
    messages = [
            {"role": "system", "content": """
    You are a text analysis expert assisting the user in identifying missing elements within a subsection of a larger text to ensure downstream consistency. You will be provided with the initial text, its subsections, a target subsection, and the currently extracted text for that subsection. Your task is to analyze these elements and determine what information is missing for the target subsection to ensure consistency. 
    Answer in a valid JSON format as follows:
    ```json
    {{
    "consistency_elements" : "concise rationale emphasizing the needed elements for consistency, if any" // think specific names, ideas or acronyms that might be lost downstream 
    
    }} 
    ```
    """}, 
        {"role": "user", "content": " {} For this particular task instance, the following elements are provided:\n Initial Content\n{}\n\n Subsections\n{}\n\n Target Subsection\n{}\n\n Current Extracted Text\n{}\n\nGiven the initial content, a list of subsections, the target subsection, and the current extracted text for that subsection, identify and return any missing elements necessary to ensure consistency.\n\n".format(icl_examples, initial_content, subsections, target_subsection, current_extracted_text)}
    ]
    
    # results = momeutils.parse_json(momeutils.ask_llm(messages, model = "g4o"))
    parseable = False
    while not parseable: 
        initial_answer, parseable, results = momeutils.safe_llm_ask(messages, model = 'g4o') 
    momeutils.crprint(json.dumps(results, indent = 4))
    return results["consistency_elements"]


def outline_report(core_idea = None, current_structure = None, user_instruction = None): 
    
    icl_examples = momeutils.load_icl(inspect.currentframe())
    
    messages = [
            {"role": "system", "content": """
    You are an expert report writer assisting the user in structuring a report given its core idea. Concretely, you are expected to propose a hierarchical outline of the report given a short text describing its core idea.
    Answer in a JSON format as follows:
    ```json
    {{
    
        "structure" : [
            ["Name section 0", 
             ["Name subsection 0.0", "Name subsection 0.1", 
                                    ["Name subsubsection 0.1.0", "Name subsubsection 0.1.1"]
             "Name subsection 0.2"],
            ["Name section 1", 
             ["Name subsection 1.0",  
                ["Name subsubsection 1.0.0", "Name subsubsection 1.0.1"], 
             ... // add more items as needed...
        ]
    }} 
    ```
    """}, 
        {"role": "user", "content": " {} For this particular task instance, the following elements are provided:\n Core Idea\n{}\n\nGiven a core idea for a report, generate a structured outline encompassing sections, subsections, and subsubsections, and return this structure under the key 'structure'.\n\n{}\n\n{}\n\nFormatting wise, Write Everything With Title (Even Words Like 'In', 'And', 'Of' ... ), avoid non alphabetic characters and refrain from using acronyms (e.g:  write Artifical Intelligence instead of AI)".format(icl_examples, core_idea, 
                                                                                                                                                                                                                                                                                                                                   "Current structure: \n{}".format(json.dumps(current_structure, indent = 4)) if current_structure else "",
                                                                                                                                                                                                                                                                                                                                   "User instruction: \n{}".format(user_instruction) if user_instruction else "")}
    ]

    
    # results = momeutils.parse_json(momeutils.ask_llm(messages, model = "g4o"))
    parseable = False
    while not parseable: 
        initial_answer, parseable, results = momeutils.safe_llm_ask(messages, model = 'g4o') 
    momeutils.crprint(json.dumps(results, indent = 4))
    return results["structure"]

def generate_section_titles(text = None, num_sections = None): 
    
    icl_examples = momeutils.load_icl(inspect.currentframe())
    
    
    messages = [
            {"role": "system", "content": """
    You are an expert editor assisting the user in structuring long texts. Concretely, you are expected to propose a list of relevant titles for subsections given the input text and the desired number of subsections.
    Answer in a JSON format as follows:
    ```json
    {{
    
        "suggested_titles" : [
            "title 0", "title 1", // add more items as needed...
        ]
    
    }} 
    ```
    """}, 
        {"role": "user", "content": " {} For this particular task instance, the following elements are provided:\n Text\n{}\n\n Num Sections\n{}\n\nGiven a text and the number of sections desired, suggest titles for each section.\n\n".format(icl_examples, text, num_sections)}
    ]
    
    # results = momeutils.parse_json(momeutils.ask_llm(messages, model = "g4o"))
    parseable = False
    while not parseable: 
        initial_answer, parseable, results = momeutils.safe_llm_ask(messages, model = 'g4o') 
    momeutils.crprint(json.dumps(results, indent = 4))
    return results["suggested_titles"]



def generate_paragraph_subcontents(structured_text = None, nb_paragraphs = None): 
    
    icl_examples = momeutils.load_icl(inspect.currentframe())
    
    
    messages = [
            {"role": "system", "content": """
You are an AI assistant expert in text analysis assisting the user in preparing accurate and efficient guidelines for textual content generation. Concretely, based on provided high-level control structure, you are expected generate {n} paragraph guidelines in a list of targeted subcontents (each element in the list will be used to produce a separate paragraph).

The provided controls can include: 
* Hierarchy location: tells you where the current paragraph fits in the broader context
* High-level context: What the current section is 
* Content: The main content of the paragraph(s) --> the message we want to convey
* Things said before: What was mentioned before in the section (if anything). This way, you can build upon more complex and sophisticated ideas and reuse the introduced concepts
* Where this is going: The following parts of the current section (if any). If provided, use those  as a reference to create a coherent and logical flow of ideas and ensure smooth transitions between paragraphs
* Number of paragraphs: The number of paragraphs you should generate for the provided contents
* Examples: Examples that the user expects for further illustrating the content (optional)
* Specific viewpoint: A specific viewpoint that the user wants to be considered (optional)
* References: References that the user wants to be included (optional)
* Visuals: Visuals that the user wants to be included (optional)

             
Answer in a JSON format as follows:
    ```json
    {{
    
        "list_of_contents" : [
{p_desc}
        ]
    
    }} 
    ```
    """.format(n = nb_paragraphs, 
               p_desc = ",\n".join(["High-level guidelines (a sentence or two) for paragraph {}".format(i) for i in range(nb_paragraphs)]))}, 
        {"role": "user", "content": " {} For this particular task instance, the following elements are provided:\n Text generation controls \n{}\n\nGiven the provided controls, produce {} grounded paragraphes guidelines under the key 'list_of_contents' that will be used for downstream generation.\n\n".format(icl_examples, structured_text, nb_paragraphs)}
    ]
    
    # results = momeutils.parse_json(momeutils.ask_llm(messages, model = "g4o"))
    parseable = False
    while not parseable: 
        initial_answer, parseable, results = momeutils.safe_llm_ask(messages, model = 'g4o') 
    momeutils.crprint(json.dumps(results, indent = 4))
       
    
    return results["list_of_contents"]

def determine_text_relevance_to_section(text = None, hierarchy = None, target_section = None): 
    
    icl_examples = momeutils.load_icl(inspect.currentframe())
    
    
    messages = [
            {"role": "system", "content": """
    You are an AI assistant with expertise in Natural Language Processing, assisting the user in organizing textual information within a hierarchical structure. Concretely, you are expected to assess the relevance of a given text snippet to a specific section, considering the provided document hierarchy. 
    Answer in a JSON format as follows:
    ```json
    {{
        "target_section_parent" : "The parent section of the target section",
        "short_rationale" :"A brief explanation of why the considered text makes sense (or not) regarding the hierarchy", 
        "score" : "A score ranging from 0 to 10 indicating the relevance of the provided text to the target section"
    
    }} 
    ```
    """}, 
        {"role": "user", "content": " {} For this particular task instance, the following elements are provided:\n Text\n{}\n\n Hierarchy\n{}\n\n Target Section\n{}\n\nGiven a text snippet, a document hierarchy, and a target section, determine the relevance of the text to the target section and return a score between 0 and 10, with 10 indicating perfect relevance.\n\n".format(icl_examples, text, hierarchy, target_section)}
    ]

    parseable = False
    while not parseable: 
        initial_answer, parseable, results = momeutils.safe_llm_ask(messages, model = 'g4o') 
    momeutils.crprint(json.dumps(results, indent = 4))
    
    return results["score"]

def paragraph_writer(paragraph_template = None): 
    
    icl_examples = momeutils.load_icl(inspect.currentframe())
    

    messages = [
            {"role": "system", "content": """
    You are a writing assistant with expertise in producing well-structured and coherent paragraphs. You are expected to use the provided template and precisely follow the guidelines. 
    Answer in a JSON format as follows:
    ```json
    {{
        "results":{{
{result}
             }}
    }} 
    ```
    """.format(result = ",\n".join(["\"paragraph_{i}\": \"Paragraph {i} based on the provided template and critical elements.\"".format(i = i) for i in range(len(paragraph_template))]))},
        {"role": "user", "content": " {} For this particular task instance, the following elements are provided:\n Paragraph Template\n{}\n\nUse the provided paragraph template to produce a well-structured and coherent text under the key 'result'.\n\n".format(icl_examples, paragraph_template)}
    ]
    momeutils.dj(messages)
    results = momeutils.parse_json(momeutils.ask_llm(messages, model = "g4o"))
    momeutils.crprint(json.dumps(results, indent = 4))
    
    return results["result"]

def paragraph_writer2(paragraph_template = None): 
    
    icl_examples = momeutils.load_icl(inspect.currentframe())
    

    messages = [
            {"role": "system", "content": """
You are a writing assistant with expertise in producing well-structured and coherent contents. You are expected to use the provided high-level layout and generate clear, well-structured and neutral sounding text for a research evaluation by an external auditing institution. 
Answer in a JSON format as follows:
    ```json
    {{
        "contents":"Your resulting contents"
    }} 
    ```
    """},
        {"role": "user", "content": " {} For this particular task instance, the following elements are provided:\n High-level layout\n{}\n\nUse the provided layout template to produce a well-structured and coherent text under the key 'contents'. To prevent formatting issues, use '\\n' in your answer to represent line breaks \n\n. ".format(icl_examples, "\n* ".join([''] + paragraph_template))}
    ]
    # Everything must be written in **french**. Tout doit être rédigé en français.
    parseable = False
    while not parseable: 
        initial_answer, parseable, results = momeutils.safe_llm_ask(messages, model = 'g4o') 
        # print(initial_answer)
    momeutils.crprint(json.dumps(results, indent = 4))
    
    return results["contents"]

def paragraph_writer_enhanced(paragraph_template = None): 

    icl_examples = momeutils.load_icl(inspect.currentframe())
    
    messages = [
            {"role": "system", "content": """
You are a writing assistant with expertise in producing well-structured and coherent paragraphs. You are expected to generate contents based on a provided template. 
Specifically, here are some controls that will inform your produced text: 
* Hierarchy location: tells you where the current paragraph fits in the broader context
* High-level context: What the current section is 
* Content: The main content of the paragraph(s) --> the message we want to convey
* Things said before: What was mentioned before in the section (if anything). This way, you can build upon more complex and sophisticated ideas and reuse the introduced concepts
* Where this is going: The following parts of the current section (if any). If provided, use those  as a reference to create a coherent and logical flow of ideas and ensure smooth transitions between paragraphs
* Number of paragraphs: The number of paragraphs you should generate for the provided contents
* Examples: Examples that the user expects for further illustrating the content (optional)
* Specific viewpoint: A specific viewpoint that the user wants to be considered (optional)
* References: References that the user wants to be included (optional)
* Visuals: Visuals that the user wants to be included (optional)
                 
Answer in a JSON format as follows:
    ```json
    {{
        "paragraphs_overlay": [what the first sub-paragraph is about, what the second sub-paragraph is about, ...], // if nb_paragraphs > 1
        "result" : "The generated paragraph based on the provided template and critical elements."
    }}
    ```
    """}, 
        {"role": "user", "content": " {} For this particular task instance, the following elements are provided:\n Paragraph Template\n{}\n\nUse the provided paragraph template to produce a well-structured and coherent paragraph under the key 'result'.\n\n".format(icl_examples, paragraph_template)}
    ]
    
    results = momeutils.parse_json(momeutils.ask_llm(messages, model = "g4o"))
    momeutils.crprint(json.dumps(results, indent = 4))
    
    return results["result"]


def simple_extract(sample_text = None, sections = None): 
    
    icl_examples = momeutils.load_icl(inspect.currentframe())
    
    messages = [
            {"role": "system", "content": """
You are an expert annotator assisting the user in structuring text. You will be provided with a text and a list of sections (for an global perspective).  
Your task is to identify and sum up how the root text could fill each target section, while being mindful of the broader context. If nothing matches, answer with 'null'
Answer in a **valid** JSON format as follows:
    ```json
    {{
{sections}
    }} 
    ```
    """.format(sections = "\n".join(['"{sectionf}": What could fill the **{section}** subsection, based on the provided context",'.format(sectionf = section.lower().strip().replace(' ', "_"), section = section) for section in sections]))}, 
     
        {"role": "user", "content": " {} For this particular task instance, the following elements are provided:\n Root text (serving as a knowledge repository)\n{}\n\n Sections\n{}\n\nGiven the root text and a list of sections, sum up some elements from the root text that could serve as justifications or supporting information for our section of interest (feel free to rephase instead of simply copy pasting things out). Try and be accurate and mindful of other sections, that is, not all content must be used and if nothing relates to the current elements, please use null.\n\n".format(icl_examples, sample_text, sections)}
    ]

    # momeutils.dj(messages)
    parseable = False
    while not parseable: 
        initial_answer, parseable, results = momeutils.safe_llm_ask(messages, model = 'g4o') 
        # print(initial_answer)
    momeutils.crprint(json.dumps(results, indent = 4))
    
    return  results


def init_config_file(config_path, **kwargs): 
    config = {
        "base_knowledge_path": os.path.join(os.path.dirname(__file__), 'sample_ainimals.txt'),
        "current_hash": None,
        "interactive_graph_path": os.path.join(os.path.dirname(__file__), 'sir_interactive_graph'),
        "tmp_structure_file": os.path.join(os.path.dirname(__file__), '.tmp_structure.json'),
        "tmp_init_text": None,
        "report_structure": None,
        "focus_node": None, 
        "last_nodes_added": [],
        "user_instruction" : None, 
        "results_path": os.path.join(os.path.dirname(__file__), 'results'),
        "report_path": os.path.join(os.path.dirname(__file__), 'reports'),
        "control_key": None, 
        "control_contents": "", 
        "review_contents": "",
        "inadequate_contents": None, 
        "non_filled": None, 
        "paragraph_parameter": None, 
        "paragraph_parameter_content": None,
        "report_details": {
            "author": "MoMe3600", 
            "title": "T-800: From the Future",
            "template": os.path.join(os.path.dirname(__file__), '.latex_template.txt')
        }, 
        "eval_follow_up_threshold": 5,
        "eval_type": None,
        "eval_results": {}
    }

    for k,v in kwargs.items():
        if k in config.keys():
            config[k] = v

    with open(config_path, 'w') as f:
        json.dump(config, f, indent = 4)

    return config


def load_config(config_path, **kwargs):
    if not os.path.exists(config_path):
        init_config_file(config_path)
    with open(config_path, 'r') as f:
        config = json.load(f)

    nb_updates = 0 
    for k,v in kwargs.items():  
        if k in config.keys():
            # print(k, v)
            nb_updates += 1
            config[k] = v

    if not "focus_node" in kwargs.keys(): 
        momeutils.crline('Loading focus node from active node')
        config['focus_node'] = mome.get_active_node(config['interactive_graph_path'])
        nb_updates += 1

    if nb_updates > 0:
        momeutils.crline('Updated config file')
    save_config(config, config_path)
    
    return json.load(open(config_path))

def save_config(config, p ): 
    with open(p, 'w') as f:
        json.dump(config, f, indent = 4)

def format_dynasty(dynasty, keep_path = False):
    # recursive function that scrolls through the dynasty (path, children) and each path is rewritten as os.path.basename(path)
    # and the children are recursively formatted
    p = os.path.basename(dynasty['path']).split('.')[0]
    formatted = {"path": dynasty['path'] if keep_path else p,
                  "name":  tmp_key_formatting(p.split('_')[-2], up_ = True) if '_' in p else p, 
                  "children": []}
    if len(dynasty['children']) > 0:
        for child in dynasty['children']:
            formatted['children'].append(format_dynasty(child, keep_path=keep_path))

    return formatted

def formatted_path_to_children(dynasty, children): 
    result= path_to_children(dynasty, children)

    root_name = result[0]
    # result[0] = 'Root'
    cleaned_result = ['Root']
    for i, r in enumerate(result[1:]):
        cleaned_result.append(r.split('_')[-2])
    return cleaned_result


def path_to_children(dynasty, target_children):

    # provided the full dynasty from the root, it returns the sequence of parents leading to the target_children
    if dynasty['path'] == target_children:
        return [dynasty['path']]
    else:
        for child in dynasty['children']:
            result = path_to_children(child, target_children)
            if result:
                return [dynasty['path']] + result



def format_section_title(s):
    return "".join([w.capitalize() for w in s.replace('_', ' ').lower().strip().split()])

def gf_all(structure, show_structure = False):

    if show_structure: 
        result = []
        if isinstance(structure, dict):
            for k in structure.keys(): 
                
                # result.append([k, 'paragraph' if all(isinstance(vv, str) for vv in structure[k]) else 'section'])
                result.append([k, 'section'])
        elif isinstance(structure, list): 
            result = [[gf(vv), 'paragraph' if isinstance(vv, str) else 'section'] for vv in structure]
        # else: 
            # result.append([k, ])
            # result.append([gf(vv), 'paragraph' if all(isinstance(vv, str) for vv in structure[k]) else 'section'])
        return result
    # else:
    #     if isinstance(structure, dict): 
    #         # momeutils.dj({'Dict': structure})
    #     elif isinstance(structure, list): 
    #         print([gf(v) for v in structure])
    #         # momeutils.dj({'List': structure})
    #     else:   
    #         # momeutils.dj({'Other': structure})
        # return [[gf(v[vv]), 'paragraph' if isinstance(vv, str) else 'section'] for vv in v]  
    # return [gf(structure) for k in structure.keys()]
    return [gf(v) for v in structure]
    # return gf(structure)

def gf(v): 
    """
    get first value, whichever type it is 
    """
    if isinstance(v, list): 
        return v[0]
    elif isinstance(v, dict):
        return list(v.keys())[0]
    elif isinstance(v, str): 
        return v
    else: 
        raise TypeError('Unrecognized type for gf --> v type: {}'.format(type(v)))




def create_nodes_recursively(root_hash, graph_folder, structure, parent_path=None, level=0, part_add=0, parent_hash=None):
    for n_idx, (key, value) in enumerate(structure.items()):

        if isinstance(value, list):
            # Create the current node

            node_name = f"lvl{level}_part{n_idx + part_add}_{format_section_title(key)}_{parent_hash}" if parent_hash else f"{parent_hash}"
            current_hash = mome.get_short_hash(node_name, 15)

            momeutils.crline('For node {} -- Value : {}'.format(node_name, value))

            node_contents = {
                # "Base contents": momeutils.j_deco(setup_base_contents_from_hls(value)),
                "Control center": momeutils.j_deco(setup_control_center_from_hls(value)),
                "Section structure": momeutils.j_deco(setup_section_structure_from_hls(value)),
                "Results": ""
            }

            # Update control center with the structure

            # Update base contents 
            # if level == 0: 
                # print('Creating node {}'.format(node_name))
                # momeutils.dj(value)
                # momeutils.dj(node_contents)
            created_node_path = mome.add_node_to_graph(
                graph_folder=graph_folder,
                contents=node_contents,
                parent_path=parent_path,
                tags=["sub_" * level + "section"],
                name_override=node_name,
                use_hash=False
            )

            # momeutils.uinput('Added {}'.format(node_name))
            

            # THESE ARE FOR AUTOMATICALLY FILLING THE PARAGRAPH CONTROLS 
            section_population = []
            # input(value)
            # section_depopulation = gf_all(value)1:]#[gf(v) for v in value][1:]
            section_depopulation = gf_all(value)[1:]
            # input(section_depopulation)
            # Recursively create child nodes
            for i, item in enumerate(value):
                
                if isinstance(item, dict):
                    create_nodes_recursively(root_hash, graph_folder, item, created_node_path, level + 1, i, current_hash)
                else:
                    # Create leaf node

                    leaf_name = f"lvl{level + 1}_part{i}_{format_section_title(item)}_{current_hash}"
                    # full_dynasty_formatted = format_dynasty(mome.collect_dynasty_paths(os.path.join(graph_folder, root_hash + ".md"), include_root=True, preserve_hierarchy=True))
                    # print(full_dynasty_formatted)
                    # print(leaf_name)
                    # input(' ok ')
                    # hierarchy_location = path_to_children(full_dynasty_formatted, leaf_name)
                    high_level_context = key # global view 
                    # print(leaf_name, key)
                    things_said_before = section_population if len(section_population) >0 else None,
                    where_this_is_going = section_depopulation if len(section_depopulation) > 0 else None

                    leaf_contents = {
                        
                        "Control center": momeutils.j_deco(initiate_control_center(
                                                                                hierarchy_location = 'to_fill',
                                                                                high_level_context = high_level_context,
                                                                               content = f"We gotta talk about {item}", 
                                                                               things_said_before = things_said_before, 
                                                                               where_this_is_going = where_this_is_going,)),
                        "Results": ""
                    }
                    mome.add_node_to_graph(
                        graph_folder=graph_folder,
                        contents=leaf_contents,
                        parent_path=created_node_path,
                        # tags=["sub_" * (level + 1) + "section"],
                        tags = ['results'],
                        name_override=leaf_name,
                        use_hash=False
                    )
                    section_population.append(item)
                    if len(section_depopulation)>0: 
                        section_depopulation.pop(0)
        else:
            # Handle other cases if necessary
            pass

def reset_config(config_path = None, **kwargs):
    if os.path.exists(config_path):
        os.remove(config_path)
    init_config_file(config_path, **kwargs)

def get_current_config(config_path = None, **kwargs):
    return momeutils.dj(load_config(config_path, **kwargs), do_break = False)

def define_structure(config, text_contents, user_instruction = None):

    if not os.path.exists(os.path.dirname(config['tmp_structure_file'])):
        os.makedirs(os.path.dirname(config['tmp_structure_file']))

    current_structure = outline_report(text_contents)
    with open(config['tmp_structure_file'], 'w') as f:
        json.dump(current_structure, f, indent = 4)
    momeutils.crline('{}'.format(json.dumps(current_structure, indent = 4)))
    subprocess.run(["code" ,'-n', config['tmp_structure_file']])
    done= False
    while not done: 

        q = momeutils.uinput("Updates (done to finish)")
        if q.lower().strip() == "done": 
            break 
        else: 
            current_structure = outline_report(text_contents, current_structure = json.load(open(config['tmp_structure_file'])), user_instruction = q)
            with open(config['tmp_structure_file'], 'w') as f:
                json.dump(current_structure, f, indent = 4)
    current_structure = json.load(open(config['tmp_structure_file']))
    converted = run_structure_conversion(current_structure)
    return converted

def sanitize_str_title(s):
    return s.replace(' AI ', ' Artficial Intelligence ').title()

def run_structure_conversion(structure): 
    """
    Converts and cleans structured representation of a report structure from nested list to dict and list
    """

    converted = []
    
    for c in convert_structure(structure):
        if len(c) > 1:
            r = {sanitize_str_title(list(c[0].keys())[0]): [c[0][list(c[0].keys())[0]]][0] + [c[i] for i in range(1, len(c))][0]}
        else:
            r = {sanitize_str_title(list(c[0].keys())[0]): [c[i][list(c[i].keys())[0]] for i in range(len(c))][0]}
        
        converted.append(r)
    converted = {list(c.keys())[0]: c[list(c.keys())[0]] for c in converted}
    # momeutils.dj(converted)
    return converted

def convert_structure(data):
    # Determine if the current level should be a list or a dictionary
    if all(isinstance(i, str) for i in data):
        return [sanitize_str_title(d) for d in data]  # If all elements are strings, return as a flat list

    result = []
    i = 0
    while i < len(data):
        item = data[i]
        if isinstance(item, str):
            if i + 1 < len(data) and isinstance(data[i + 1], list):
                # Create a dictionary entry if the next item is a list
                result.append({item: convert_structure(data[i + 1])})
                i += 1  # Skip the next item since it's already processed
            else:
                # Add the string directly to the result list
                result.append(item)
        else:
            # Handle case where item is a list (shouldn't occur in this structure)
            result.append(convert_structure(item))
        i += 1

    return result





def initialize_graph_from_text(config_path=None, **kwargs):
    config = init_config_file(config_path, **kwargs)

    text_contents = open(kwargs.get('text_contents', None)).read()
    config['tmp_init_text'] = kwargs.get('text_contents', None)

    if text_contents is None:
        raise ValueError('Empty text contents')
    
    structure = define_structure(config, text_contents, kwargs.get('user_instruction', None))

    config['report_structure'] = structure
    save_config(config, config_path)

    make_graph(config_path, **kwargs)
    return 

def make_graph(config_path=None, **kwargs):
    config = load_config(config_path, **kwargs)
    mome.init_obsidian_vault(config['interactive_graph_path'], exists_ok=False)

    # CLEANING RESULTS FOLDER TO AVOID CONFLICTS
    if os.path.exists(config['results_path']):
        shutil.rmtree(config['results_path'])

    if config['report_structure'] is None:  
        # TMP STRUCTURE (OR STARTING POINT)
        structure = {
        "Operation Topic": ['Overview',
            {'Limitations Of Our Previous Perspective And What Is Wrong With Other Methods': [
                    "Focus On Local Problems Instead Of Systemic Solutions",
                    "Automated Processing Of Unstructured Data",
                    "Knowledge Management"
            ]
            },
            "What We Did"
        ],
        "Technical Justification": [
            "Test"
        ],
        "What We Actually Did": [
            "Test"
        ],
        "Conclusion": [
            "Test"
        ]
        }

        config['report_structure'] = structure
    else: 
        structure = config['report_structure']

    
    
    # Create root node
    root_contents = {"Control center": momeutils.j_deco(setup_control_center_from_hls(structure)),
                     "Section structure": momeutils.j_deco(setup_section_structure_from_hls(structure)),
                     "Results": ""}
    

    # Root id
    base_hash = mome.get_short_hash("hello", 15)  # TMP

    if config['tmp_init_text'] is not None:
        base_text = open(config['tmp_init_text']).read()
        
        current_root_structure = momeutils.parse_json(root_contents['Section structure'])
        current_root_structure['initial_contents'] = base_text
        root_contents['Section structure'] = momeutils.j_deco(current_root_structure)

        base_hash = mome.get_short_hash(base_text, 15)
    

    root_node_path = mome.add_node_to_graph(
        graph_folder=config['interactive_graph_path'],
        contents=root_contents,
        parent_path=None,
        tags=["root"],
        name_override=base_hash,
        use_hash=False
    )


    config['current_hash'] = base_hash
    config['results_path'] = os.path.join(os.path.dirname(config['report_path']), "results", f"results_{base_hash}.json")

    # Create nodes recursively
    create_nodes_recursively(config['current_hash'], config['interactive_graph_path'], structure, root_node_path, level=0, part_add=0, parent_hash=mome.get_short_hash(base_hash, 15))
    save_config(config, config_path)

def get_control_center(node): 
    return momeutils.parse_json(mome.get_node_section(node, "Control center"))  

def get_section_structure(node):
    return momeutils.parse_json(mome.get_node_section(node, "Section structure"))

def model_suggestion_to_user_instruction(config_path = None, **kwargs): 
    config = load_config(config_path, **kwargs)
    focus_node = get_focus_node(config)
    control_center = focus_node['control']
    structure = focus_node['structure']
    control_key = config['control_key']
    ui = config['user_instruction']

    key_id = check_matching_key(control_center, control_key)

    model_suggestion = "\n* ".join([""] + control_center[key_id]['model_suggestion'])
    p = "In the process of preparing section contents for a report, the initially provided contents didn't match the target section '{}' and some suggestions were provided, which we now want to rewrite and format to convey the high-level ideas for that section. \rInitially provided text: '{}'\n\nSuggestions: \n{}\n\nUser additional instruction: {}\n\nConsidering the broader context and the (optional) user directive, rewrite the suggestions in one or two insightful sentences.".format(key_id, 
                                                                                                                                                                                                                                                                                                                                                                structure['initial_contents'], 
                                                                                                                                                                                                                                                                                                                                                                model_suggestion, 
                                                                                                                                                                                                                                                                                                                                                                config['user_instruction']
                                                                                                                                                                                                                                                                                                                                                                )
    out = momeutils.basic_task(p, model = 'g4o')
    control_center[key_id]['user_instruction'] = out
    control_center[key_id]['model_suggestion'] = None

    # Updating 
    mome.update_section(focus_node['path'], 'Control center', momeutils.j_deco(control_center))
    save_config(config, config_path)


def sub_inspiration(config_path = None, **kwargs):
    config = load_config(config_path, **kwargs)
    focus_node = get_focus_node(config)
    structure = focus_node['structure']

    # if "nb_subs" in kwargs.keys():
    #     nb_subs = int(kwargs['nb_subs'])
    nb_subs = kwargs.get('nb_subs', 3)
    # p = "Given the following existing contents, suggest {} titles for children parts. Be careful to avoid special characters.\n\n Existing contents: {}".format(nb_subs, structure['initial_contents'])
    out = generate_section_titles(structure['initial_contents'], nb_subs)

    structure['subs_titles'] = out
    mome.update_section(focus_node['path'], 'Section structure', momeutils.j_deco(structure))
    save_config(config, config_path)
    



def propagate(config_path = None, **kwargs): 

    config = load_config(config_path, **kwargs)
    focus_node = get_focus_node(config)

    # THOSE ENABLE LAST MINUTE CHANGES TO THE STRUCTURE 
    check_children_node_existence(config)
    update_children_node_state(config)

    inadequate_contents, non_filled = do_propagate(config_path, **kwargs)
    config = load_config(config_path, **kwargs)
    momeutils.crline("{}".format(json.dumps({"inadequate_contents": inadequate_contents, "non_filled": non_filled}, indent = 4)))
    config['inadequate_contents'] = inadequate_contents
    config['non_filled'] = non_filled
    save_config(config, config_path)

def do_propagate(config_path=None, **kwargs):
    config = load_config(config_path, **kwargs)
    full_dynasty_formatted = format_dynasty(mome.collect_dynasty_paths(os.path.join(config['interactive_graph_path'], config['current_hash'] + ".md"), include_root=True, preserve_hierarchy=True))
    focus_node_path = os.path.join(config['interactive_graph_path'], config['focus_node'] + ".md")
    control_center = get_control_center(focus_node_path)
    links = mome.get_node_links(focus_node_path)
    inadequate = {}
    non_filled = {}
    for i, (k,l) in enumerate(zip(control_center.keys(), links)):
        link_path = os.path.join(config['interactive_graph_path'], l + ".md")
        if control_center[k]['template'] == "direct": 
            link_control_center = get_control_center(os.path.join(config['interactive_graph_path'], l + ".md"))

            before = list(control_center.keys())[:i]
            before = before if len(before) > 0 else None
            after = list(control_center.keys())[i+1:]
            after = after if len(after) > 0 else None

            
            hierarchy_location = formatted_path_to_children(full_dynasty_formatted, l)

            link_control_center = update_existing_structure(link_control_center, 
                                                            hierarchy_location = hierarchy_location,
                                                            content = control_center[k]['user_instruction'], 
                                                            things_said_before = before,
                                                            where_this_is_going = after)
            mome.update_section(link_path, 
                                "Control center",
                                momeutils.j_deco(link_control_center))
        elif control_center[k]['template'] == "default": 
            # insert in initial contents

            momeutils.crline('Updating section structure in {}'.format(l))
            section_structure = momeutils.parse_json(mome.get_node_section(link_path, "Section structure"))
            section_structure = update_existing_structure(section_structure, 
                                                          initial_contents = control_center[k]['user_instruction'])
            mome.update_section(link_path,
                                "Section structure", 
                                momeutils.j_deco(section_structure))
            if not section_structure['subs_titles'] == ['To Fill']:
                leftovers = pour_info(config_path, focus_node = l)
                # momeutils.dj({"Focus node": config['focus_node'], 
                #               "L": l, 
                #               "Leftovers": leftovers})
                # inadequate[config['focus_node']] = leftovers
                inadequate[l] = leftovers
                # momeutils.dj({"leftovers": leftovers, 
                #             "inadequate": inadequate})
                # inadequate[l] = do_propagate(config_path, focus_node = l)
                prop_results = do_propagate(config_path, focus_node = l)
                inadequate[l] += prop_results[0]
                non_filled[l] = prop_results[1]
            else: 
                non_filled[l] = 'to_fill'
    return inadequate, non_filled

def collect_hierarchy_to_focus_node(config):
    return collect_hierarchy_to_children(config, config['focus_node'])
     
def collect_hierarchy_to_children(config, target_node_path, raw = False): 
    root_path = os.path.join(config['interactive_graph_path'], config['current_hash'] + ".md")
    full_dynasty = mome.collect_dynasty_paths(root_path, include_root=True, preserve_hierarchy=True)
    formatted_full = format_dynasty(full_dynasty)
    if raw:
        to_children = path_to_children(formatted_full, target_node_path)
    else:
        to_children = formatted_path_to_children(formatted_full, target_node_path)
    # hierarchy = formatted_path_to_children(format_dynasty(mome.collect_dynasty_paths(root_path, include_root=True, preserve_hierarchy=True)), target_node_path)
    return to_children


def add_default_hierarchical_node(focus_node_contents, node_name, parent_path = None): 
    control_center = focus_node_contents['control']
    lvl, part, name, hash_ = split_node_for_info(focus_node_contents['path'])
    
    
    current_node_key = tmp_key_formatting(node_name.split('_')[-2], up_ = True)
    new_control_center = setup_control_center([["To Fill", control_center[current_node_key]['template']]])
    new_section_structure = setup_section_structure(initial_contents = "Empty", subs_titles = ['To Fill'])
    node_contents = {"Control center": momeutils.j_deco(new_control_center), 
                    "Section structure": momeutils.j_deco(new_section_structure), 
                    "Results": ""} 
    tags = ["sub_"*(lvl+1) + "section"]

    return mome.add_node_to_graph(
        graph_folder=os.path.dirname(focus_node_contents['path']),
        contents=node_contents,
        parent_path=parent_path, # putting it to None and adding it later otherwise it gets added at the end 
        tags = tags, 
        name_override= node_name,
        use_hash=False
    )

def add_default_result_node(focus_node_contents, node_name, parent_path = None): 
    new_control_center = initiate_control_center()
    node_contents = {"Control center": momeutils.j_deco(new_control_center), 
                    "Results": ""}
    tags = ['results']

    return mome.add_node_to_graph(
        graph_folder=os.path.dirname(focus_node_contents['path']),
        contents=node_contents,
        parent_path=parent_path,
        tags=tags,
        name_override=node_name,
        use_hash=False
    )



def check_children_node_existence(config): 

    focus_node_contents = get_focus_node(config)
    control_center = focus_node_contents['control']
    lvl, part, name, hash_ = split_node_for_info(focus_node_contents['path'])
    focus_node_hash = mome.get_short_hash(momeutils.bn(focus_node_contents['path']), 15)
    if is_root(config, focus_node_contents['path']):
        lvl = -1
    
    

    # print(momeutils.bn(focus_node_contents['path']))
    # print(lvl, part, name, hash_)
    # input('ok ? ')
    # Constructing node name and creating node if currently missing 
    for i,k in enumerate(control_center.keys()): 
        node_name = f"lvl{lvl+1}_part{i}_{k.replace(' ', '')}_{focus_node_hash}"
        node_path = os.path.join(os.path.dirname(focus_node_contents['path']), node_name + ".md")

        add_children_paragraph = False 
        # momeutils.crline(node_name)
        if not os.path.exists(node_path):
            # input('missing {}'.format(node_name))
            if control_center[k]['template'].strip() == "default": 
              
                # input('about to hierarchical add {}'.format(node_name))
                added_node = add_default_hierarchical_node(focus_node_contents, node_name)
                add_children_paragraph = True 
            else:
                # input('about to result add {}'.format(node_name))
                added_node = add_default_result_node(focus_node_contents, node_name)
         

            # Since the node is added, assume it is user validated --> Add it to links 
            mome.add_link_at_position(focus_node_contents['path'], node_name , i)

            # IF THE NODE IS A SECTION (HIERARCHY) NODE, 
            # WE AUTOMATICALLY CREATE A PLACEHOLDER CHILD TO ENSURE CORRECT DOWNSTREAM USAGE OF STRUCTURE
            if add_children_paragraph: 
                new_section_structure = get_section_structure(added_node)
                section_hash = mome.get_short_hash(node_name, 15)
        
                p_name = f"lvl{lvl+2}_part0_{new_section_structure['subs_titles'][0].replace(' ', '')}_{section_hash}"
            
                # input('Following hierachical add, adding result node {}'.format(p_name))
                add_default_result_node(focus_node_contents, p_name, parent_path = added_node)

                current_dynasty = format_dynasty((mome.collect_dynasty_paths(focus_node_contents['path'], include_root=True, preserve_hierarchy=True)), keep_path= True)
                node_index = [i for i, c in enumerate(current_dynasty['children']) if c['path'] == node_path][0]
                current_dynasty['children'][node_index] = {"path": node_path, "children": [{"path": os.path.join(os.path.dirname(node_path), p_name + ".md"), "children": []}]}
                config['report_structure'] = update_report_structure(config, current_dynasty, focus_node_contents['path'])
    return config


def update_children_node_state(config):
    """
    For each children, checks if the template matches 
    If there are changes, 
        * Removes the dynasty, 
        * Adds the relevant node (ideally, node contents should be somehow kept )
        * Updates the config report structure accordingly
    """
    focus_node_contents = get_focus_node(config)
    control_center = focus_node_contents['control']
    lvl, part, name, hash_ = split_node_for_info(focus_node_contents['path'])
    if is_root(config, focus_node_contents['path']):
        lvl = -1

    focus_node_hash = mome.get_short_hash(momeutils.bn(focus_node_contents['path']), 15)
    for i, k in enumerate(control_center.keys()): 
        # IS THERE DISCREPANCY BETWEEN TEMPLATE AND ACTUAL NODE TYPE
        # VARIOUS WAYS TO CHECK -> FIRST APPROXIMATION USING TEMPLATE SECTIONS 
        user_request = control_center[k]['template'].strip().lower()
        node_path = os.path.join(os.path.dirname(focus_node_contents['path']), f"lvl{lvl+1}_part{i}_{k.replace(' ', '')}_{focus_node_hash}.md")
        node_type = 'default' if mome.check_section(node_path, "Section structure") else 'direct'
        
        # ADDING THE NODE 
        if node_type != user_request: 
            clean_dynasty(node_path, include_root = True)
            if user_request == "direct": 
                child_node_name = f"lvl{lvl+1}_part{i}_{k.replace(' ', '')}_{focus_node_hash}"  
                add_default_result_node(focus_node_contents, child_node_name)
            elif user_request == "default":
                child_node_name = f"lvl{lvl+1}_part{i}_{k.replace(' ', '')}_{focus_node_hash}"
                add_default_hierarchical_node(focus_node_contents, child_node_name)
            else: 
                raise ValueError(f'Unrecognized template type {user_request}')

    return config


def get_node_compilation_results(node_path):
    
    return True, mome.get_node_section(node_path, "Results")


def tmp_xps_eval(contents):
    p = "You are tasked with evaluating if and how convincing the scientific experiment in the provided contents are. Return a score between 0 and 10, with 10 being the most convincing. If there are no clearly defined experiments, return 0. \n\n Provided contents: {}".format(contents)
    return int(momeutils.basic_task(p, model = 'g4o'))
def tmp_check_innov(contents): 
    p = "You are tasked with evaluating the level of explicitely constructed innovation in the provided contents. Return a score between 0 and 10, with 10 being representative of highly innovative work realized. If there is no clearly defined innovation realized in the work, return 0. \n\n Provided contents: {}".format(contents)
    return int(momeutils.basic_task(p, model = 'g4o'))


def find_parent(config, node_name, hierarchy = None):
    """
    Finds the parent of a given node name using the provided hierarchy.

    :param config: Configuration dictionary containing report structure and paths.
    :param node_name: The name of the node for which to find the parent.
    :param hierarchy: The hierarchy list of the node.
    :return: The path of the parent node if found, otherwise None.
    """
    lvl, part, name, hash_ = split_node_for_info(node_name)
    if lvl == 0: 
        return os.path.join(config['interactive_graph_path'], config['current_hash'] + ".md")
    
    
    if hierarchy is None:
        hierarchy = collect_hierarchy_to_children(config, node_name)

    momeutils.dj(hierarchy)
    parent_result = None
    for i in range(10):
        parent_name = hierarchy[-2]
        parent_result = find_node_in_structure(
            config['report_structure'],
            lvl - 1,
            i,
            tmp_key_formatting(parent_name, up_=True).strip(),
            current_lvl=-1
        )
        if parent_result is not None:
            break

    if parent_result is None:
        return None

    root_path = os.path.join(config['interactive_graph_path'], config['current_hash'] + ".md")
    current_dynasty = format_dynasty(
        mome.collect_dynasty_paths(root_path, include_root=True, preserve_hierarchy=True)
    )['children']

    for cd in current_dynasty:
        if cd['name'] == parent_result[0]:
            current_dynasty = cd
            break

    

    if len(parent_result) == 1:
        return os.path.join(config['interactive_graph_path'], current_dynasty['path'] + '.md')
    # momeutils.crline(node_name)
    # momeutils.crline("{}".format(json.dumps(current_dynasty, indent=4)))
    # momeutils.crline(parent_result)


    flat_hierarchy = flatten_hierarchy(current_dynasty)
    parent_path = find_parent_path(flat_hierarchy, parent_result[-1])

    return os.path.join(config['interactive_graph_path'], parent_path.split('/')[-1] + ".md")
    

    # for pr in parent_result[1:]:
    #     if isinstance(pr, int):
    #         current_dynasty = current_dynasty['children'][pr]
    #     elif isinstance(pr, str):
    #         if current_dynasty['name'] == pr:
    #             momeutils.crline(f'Looking for {node_name} - Parent name {pr} ')
    #             input(current_dynasty)
    #             return os.path.join(config['interactive_graph_path'], current_dynasty['path'] + '.md')
    #         target_idx = [i for i, c in enumerate(current_dynasty['children']) if c['name'] == pr][0]
    #         current_dynasty = current_dynasty['children'][target_idx]

    # # return None
    # collected_path = flatten_hierarchy_and_find_path(current_dynasty, parent_result)
    # momeutils.uinput("{}".format(collected_path))
    # return os.path.join(config['interactive_graph_path'], collected_path + '.md') 

def flatten_hierarchy(hierarchy):
    """
    Flattens a nested hierarchy into a list of nodes with their paths.

    :param hierarchy: The nested hierarchy to flatten.
    :return: A list of tuples containing node names and their paths.
    """
    flat_list = []

    def _flatten(current_node, current_path):
        flat_list.append((current_node['name'], current_path))
        for child in current_node.get('children', []):
            _flatten(child, os.path.join(current_path, child['path']))

    _flatten(hierarchy, hierarchy['path'])
    return flat_list

def find_parent_path(flat_hierarchy, target_name):
    """
    Finds the path of the target node in the flattened hierarchy.

    :param flat_hierarchy: The flattened hierarchy list.
    :param target_name: The name of the target node.
    :return: The path of the target node if found, otherwise None.
    """
    for name, path in flat_hierarchy:
        if name == target_name:
            return path
    return None

def find_neighbors(config, node_name, hierarchy = None):

    parent = find_parent(config, node_name, hierarchy)
    if parent is None:
        return []
    hierarchy =mome.collect_dynasty_paths(parent, include_root=True, preserve_hierarchy=True)
    neighbors = [[i, momeutils.bn(c['path'])] for i, c in enumerate(hierarchy['children']) if momeutils.bn(c['path']) != node_name]
    return neighbors



# def flatten_hierarchy_and_find_path(current_dynasty, parent_result):
#     """
#     Flattens the hierarchy and finds the path for the given parent result.

#     :param current_dynasty: The current hierarchy structure.
#     :param parent_result: The result containing the path to the parent node.
#     :return: The path of the parent node if found, otherwise None.
#     """
#     def search_hierarchy(dynasty, target):
#         if dynasty['name'] == target:
#             return dynasty['path']
#         for child in dynasty.get('children', []):
#             result = search_hierarchy(child, target)
#             if result:
#                 return result
#         return None

#     for pr in parent_result[1:]:
#         if isinstance(pr, int):
#             current_dynasty = current_dynasty['children'][pr]
#         elif isinstance(pr, str):
#             path = search_hierarchy(current_dynasty, pr)
#             if path:
#                 return path, #os.path.join(config['interactive_graph_path'], path + '.md')

#     return None


def check_missing_children(config_path=None, **kwargs):
    config = load_config(config_path, **kwargs)
    focus_node = get_focus_node(config)

    if is_leaf(focus_node['path']):
        all_leaves = [focus_node['path']]
    else:
        all_leaves = mome.collect_leaf_paths(focus_node['path'])

    missing_results = {}
    for leaf in all_leaves:
        name = momeutils.bn(leaf).split('_')[-2]
        if name == "ToFill" or get_control_center(leaf)['content'] is None:
            leaf_hierarchy = collect_hierarchy_to_children(config, momeutils.bn(leaf))
            missing_results[momeutils.bn(leaf)] = {"status": "empty", "hierarchy": leaf_hierarchy}

    momeutils.crline('Missing: \n\n{}'.format(
        json.dumps(missing_results, indent=4)
    ))
    
    
    collected_parents = []
    for k in missing_results:
        parent_path = find_parent(config, k, None) #missing_results[k]['hierarchy'])
        if parent_path:
            collected_parents.append([momeutils.bn(parent_path), momeutils.bn(k)])
            parent_controls = get_control_center(parent_path)
            parent_structure = get_section_structure(parent_path)

    momeutils.dj(collected_parents)


def update_report_structure_from_graph(config_path= None, **kwargs): 
    config = load_config(config_path, **kwargs)
    root = os.path.join(config['interactive_graph_path'], config['current_hash'] + ".md")
    structure = format_dynasty(mome.collect_dynasty_paths(root, include_root=True, preserve_hierarchy=True))
    
    raise NotImplementedError('Not implemented yet')

def manual_control_key_change(config_path = None, **kwargs):
    config = load_config(config_path, **kwargs)
    control_key = config['control_key']
    instructions = config['user_instruction']
    focus_node = get_focus_node(config)
    # get all compilable children
    targets = find_compilable_from_focus(focus_node['path'])

    for t in targets:
        control_center = momeutils.parse_json(get_control_center(t))

        if not control_key in control_center.keys():
            momeutils.crline('Control key {} not found in {}'.format(control_key, momeutils.bn(t)))
            continue
        control_center[control_key] = instructions
        # for k in control_center.keys():
        #     control_center[k]['template'] = momeutils.uinput('New template for {}'.format(k))
        mome.update_section(t, 'Control center', momeutils.j_deco(control_center)) 

    save_config(config, config_path)

def eval_follow_up(config_path = None, **kwargs):
    config = load_config(config_path, **kwargs)
    focus_node = get_focus_node(config)
    eval_type = config['eval_type']
    if not eval_type in config['eval_results'].keys():
        momeutils.crline('No evaluation results to follow up on')
        return 

    # Collecting nodes with low evaluation results
    targets = [t for t in config['eval_results'][eval_type].keys() if config['eval_results'][eval_type][t] < config['eval_follow_up_threshold']]
    if len(targets) == 0:
        momeutils.crline('All set')
        return

    for t in targets:
        momeutils.crline("Need to compose a helpful workflow to handle the following: \n{}".format(json.dumps({"Target": t, 
                      "Result": config['eval_results'][eval_type][t], 
                      "Control": get_control_center(os.path.join(config['interactive_graph_path'], t + ".md")), 
                      "parent": find_parent(config, t, None), 
                      "neighbors": find_neighbors(config, t, None)}, indent = 4)))
    
    
def find_compilable_from_focus(focus_node_path):
    if is_leaf(focus_node_path):
        return [focus_node_path]
    else:
        return collect_all_compilable_children(focus_node_path) 

def run_eval(config_path= None, **kwargs):

    config = load_config(config_path, **kwargs)
    focus_node = get_focus_node(config)
    control_center = focus_node['control']
    structure = focus_node['structure']
    eval_type = config['eval_type']
    
    # Collecting target nodes
    # if is_leaf(focus_node['path']): 
    #     targets = [focus_node['path']]
    # else: 
    #     targets = collect_all_compilable_children(focus_node['path'])
    targets = find_compilable_from_focus(focus_node['path'])
    
    if eval_type == "xps": 
        eval_func = tmp_xps_eval
    elif eval_type == "length": 
        eval_func = lambda x: len(x)
    elif eval_type == "check_innov": 
        eval_func = tmp_check_innov

    for t in targets: 
        valid, contents = get_node_compilation_results(t)

        if valid: 
            result = eval_func(contents)
        else: 
            result = None
        if not eval_type in config['eval_results'].keys():
            config['eval_results'][eval_type] = {}
        config['eval_results'][eval_type][momeutils.bn(t)] = result

    save_config(config, config_path)

def get_structure_children(node_path, config= None): 
    if config is not None: 
        node_path = os.path.join(config['interactive_graph_path'], momeutils.bn(node_path).split('.')[0] + ".md")
    
    structure = get_section_structure(node_path)
    return structure['subs_titles']

def pull_details(config_path = None, **kwargs):
    config = load_config(config_path, **kwargs)
    focus_node = get_focus_node(config)
    # control_center = focus_node['control']
    # structure = focus_node['structure']

    node_parent= find_parent(config, momeutils.bn(focus_node['path']))
    lvl, part, name, hash_ = split_node_for_info(focus_node['path'])
    parent_controls = get_control_center(node_parent)
    details = parent_controls[list(parent_controls.keys())[part]]['additional_details']
    consistency_elements = parent_controls[list(parent_controls.keys())[part]]['consistency_elements']

    out = momeutils.unconstrained_task("We are writing a report and focusing on the section '{}'. The following details were provided by the parent section {}: \n\n{}\n\nRate (0-10) the relevance of each piece of detail so we can identify the most relevant to include downstream.".format(name, 
                                                                                                                                                                                                                                                                                        split_node_for_info(momeutils.bn(node_parent))[-2],
                                                                                                                                                                                                                                                                                        details))

    momeutils.dj(out)

def look_ahead(config_path = None, **kwargs):

    config = load_config(config_path, **kwargs)
    focus_node = get_focus_node(config)
    if is_leaf(focus_node['path']): 
        momeutils.crline('Leaf node, nothing to do')
        return
    
    # Collecting all direct children nodes 
    current_hash = mome.get_short_hash(momeutils.bn(focus_node['path']), 15)
    if is_root(config, focus_node['path']):
        target_lvl = -1
    else: 
        target_lvl = int(momeutils.bn(focus_node['path']).split('_')[0].replace('lvl', ''))

    target_nodes = ["lvl{}_part{}_{}_{}".format(target_lvl+1, i, k.replace(' ', ''), current_hash) for i, k in enumerate(focus_node['control'].keys())]

    results = {}

    # compilation_results = {}

    # # PARALLELIZING THE COMPILATION
    look_ahead_params = [
        (run_look_ahead, (config, focus_node, key, tn)) for key, tn in zip(focus_node['control'].keys(), target_nodes)
    ]


    # saving focus node in case shit happens 
    shutil.copy(focus_node['path'], os.path.join(os.path.dirname(__file__), '.tmp_focus_node.md'))
    
    try:
        parallel_results = momeutils.mapper(look_ahead_params)
        
        # Applying results
        control_center = focus_node['control']
        for r, key in zip(parallel_results, control_center.keys()): 
            if isinstance(r, dict): 
                control_center[key]['look_ahead_results'] = {c:int(s) for c,s in zip(r['children'], r['scores'])}
            else: 
                control_center[key]['look_ahead_results'] = r                

        mome.update_section(focus_node['path'], 'Control center', momeutils.j_deco(control_center))


        os.remove(os.path.join(os.path.dirname(__file__), '.tmp_focus_node.md'))
        save_config(config, config_path)
    except Exception as e: 
        momeutils.crline(f'Look ahead failed (probably parallel mapping)\nError: {e}')
        shutil.move(os.path.join(os.path.dirname(__file__), '.tmp_focus_node.md'), focus_node['path'])
    
    


def run_look_ahead(config, focus_node, key, target_node):

    if is_leaf(os.path.join(config['interactive_graph_path'], target_node + ".md")):
        return "leaf"
    else: 
        future_contents = get_control_center(focus_node['path'])[key]['user_instruction'] 
        future_children = get_structure_children(target_node, config)
        out = assess_text_sufficiency(future_contents, future_children)
        # process_text_sufficiency(config, focus_node, key, future_children, out)
        return {"scores": out, "children": future_children}

  

# def process_text_sufficiency(config, focus_node, key, future_children, out):

#     control_center = get_control_center(get_focus_node(config)['path'])
    
#     # enhancing section 
#     if not 'look_ahead_results' in control_center[key].keys():
#                 control_center[key]['look_ahead_results'] = {}

#     for i, oo in enumerate(out): 
#         if oo < 5:         
#             control_center[key]['look_ahead_results'][future_children[i]] = "You messed up, dummy. Relevance score is {}".upper().format(oo)
#         else: 
#             control_center[key]['look_ahead_results'][future_children[i]] = None
    
#     mome.update_section(focus_node['path'], 'Control center', momeutils.j_deco(control_center))

def run_test(config_path = None, **kwargs): 
    config = load_config(config_path, **kwargs)
    focus_node= get_focus_node(config)

    key_concepts = extract_key_concepts(focus_node['structure']['initial_contents'])
    complementary_info = extract_complementary_info_concepts(focus_node['structure']['initial_contents'], key_concepts)

    
    
    # testing consistentcy 
    # results = simple_extract(sample_text = focus_node['structure']['initial_contents'],
    #                             sections = focus_node['structure']['subs_titles'])
    results = momeutils.parse_json(focus_node['control'])

    for k in results.keys():

        relatedness = rate_key_concepts_relatedness(focus_node['structure']['initial_contents'], 
                                                focus_node['structure']['subs_titles'], 
                                                k,
                                                results[k], 
                                                key_concepts)
        momeutils.dj(relatedness)
        # key_concepts = identify_missing_elements_for_consistency(focus_node['structure']['initial_contents'], 
        #                                                          focus_node['structure']['subs_titles'], 
        #                                                          k, results[k])
        # momeutils.uinput("")

  
def pour_info(config_path=None, **kwargs):
    config = load_config(config_path, **kwargs)

    focus_node = get_focus_node(config)
    focus_node_path = focus_node['path']

    hierarchy = collect_hierarchy_to_focus_node(config)
    section_structure = focus_node['structure']
    control_center = focus_node['control']

    subs_titles_syntax_check(focus_node)  
       
    # Ensuring node existence and managing report_structure
    config = check_children_node_existence(config)
    config = update_children_node_state(config)
    
    full_hierarchy = format_dynasty(mome.collect_dynasty_paths(os.path.join(config['interactive_graph_path'], config['current_hash'] + ".md")
                                                , include_root = True, preserve_hierarchy = True))
    # test = determine_text_relevance_to_section(text = section_structure['initial_contents'],
    #                                            hierarchy=json.dumps(full_hierarchy, indent = 4), 
    #                                            target_section=os.path.basename(focus_node_path).split('.')[0])
    

    leftovers = []
    results = simple_extract(sample_text = section_structure['initial_contents'], 
                             sections = section_structure['subs_titles'],)

    
    # Fillings for consistency and additional info
    key_concepts = extract_key_concepts(section_structure['initial_contents'])
    complementary_info = extract_complementary_info_concepts(section_structure['initial_contents'], key_concepts)
    
    for i, (k,v) in enumerate(results.items()):
        if v is None or v.strip().lower() == "null": 
            leftovers.append(k)
        control_center[list(control_center.keys())[i]]['user_instruction'] = v
    
    pour_info_params = [
        (run_pour_info, (config, control_center, section_structure, leftovers, key, results[key_r], key_concepts, complementary_info)) for key, key_r in zip(control_center.keys(), results.keys())
    ]
    parallel_results = momeutils.mapper(pour_info_params)   

    for i, (p, k)  in enumerate(zip(parallel_results, control_center.keys())): 
        for kk in p.keys(): 
            control_center[k][kk] = p[kk]
        # control_center[k][p[0]] = p[1]

    mome.update_section(focus_node_path, "Control center", momeutils.j_deco(control_center))
    save_config(config, config_path)
    return leftovers

def run_pour_info(config, control_center, section_structure, leftovers, key, results, key_concepts = None, complementary_info = None):

    if key in leftovers: 
        result = "Some model suggestion"
        return ["model_suggestion", result]
    else: 
        if key_concepts is not None:
            ratings= rate_key_concepts_relatedness(section_structure['initial_contents'], section_structure['subs_titles'], key, results, key_concepts)
            result = [key_concepts[i] for i,(k,v) in enumerate(ratings.items()) if v > 4]
        else: 
            result = "Check consistency " #identify_missing_elements_for_consistency(section_structure['initial_contents'], section_structure['subs_titles'], key, control_center[key]['user_instruction'])
        
        if complementary_info is not None: 
            ci = complementary_info
        else: 
            ci = "Check additional info"
        # return ["consistency_elements", result]
        return {"consistency_elements": result, "additional_details": ci}


    # for i, (s,k) in enumerate(zip(section_structure['subs_titles'], results.keys())): 
    #     control_center[s]['user_instruction'] = results[k]
        
    #     if k in leftovers: 
    #         # TODO add something to help enhancing that specific key 
    #         control_center[s]['model_suggestion'] = "Some model suggestion"
    #     # Add additional details here
    #     else: 
    #         # Adding consistency details to maintain coherence 
    #         out = identify_missing_elements_for_consistency(section_structure['initial_contents'], section_structure['subs_titles'], k, control_center[s]['user_instruction'])
    #         control_center[s]['consistency_elements'] = out
    
    # # updating sections 
    # mome.update_section(focus_node_path, "Control center", momeutils.j_deco(control_center))

    # save_config(config, config_path)
    # return leftovers

def model_inspiration(config_path = None, **kwargs):
    config = load_config(config_path, **kwargs)
    focus_node = get_focus_node(config)
    # ASSUMING WE FOCUS ON A KEY IN THE CONTROL CENTER 
    control_center = focus_node['control']
    structure = focus_node['structure']
    control_key = config['control_key']
    ui = "Specific user guiding: {}".format(config['user_instruction']) if config['user_instruction'].strip() != "" else ""

    lvl, part, name, _ = split_node_for_info(focus_node['path'])

    key_id = check_matching_key(control_center, control_key)
    current_contents = control_center[key_id]['user_instruction']
    
    momeutils.unconstrained_task(f"In the context of writing a report, we're currently focused on the section '{name}' with the following contents\n\n{json.dumps(control_center, indent = 4)}\n\n The user is lacking inspiration for the part named {key_id}, which currently contains the following: {current_contents}\n\n{ui}\n\nProvide some inspiration to enhance and develop the target section {key_id}", model = "g4o")  

def set_templates(config_path= None, **kwargs):
    config = load_config(config_path, **kwargs)
    focus_node = get_focus_node(config)
    control_center = focus_node['control']

    assert 'templates' in kwargs.keys(), "No templates provided in config"
    assert len(kwargs['templates']) == len(control_center.keys())
    for i, k in enumerate(control_center.keys()):
        control_center[k]['template'] = kwargs['templates'][i].strip()
    mome.update_section(focus_node['path'], "Control center", momeutils.j_deco(control_center))
    save_config(config, config_path)

def direct_structure_update(config_path= None, **kwargs):   
    config = load_config(config_path, **kwargs)
    config = check_children_node_existence(config)
    config = update_children_node_state(config)
    save_config(config, config_path)


def set_subtitles(config_path= None, **kwargs):
    config = load_config(config_path, **kwargs)
    focus_node = get_focus_node(config)
    structure = focus_node['structure']

    assert 'subtitles' in kwargs.keys(), "No subtitles provided in config"

    structure['subs_titles'] = [k.strip().title() for k in kwargs['subtitles']]
    mome.update_section(focus_node['path'], "Section structure", momeutils.j_deco(structure))
    save_config(config, config_path)

def expand_initial_contents(config_path = None, **kwargs):
    config = load_config(config_path, **kwargs)
    focus_node = get_focus_node(config)
    control_center = focus_node['control']


    if is_leaf(focus_node['path']):
        target_content = control_center['content']
        structure = {"initial_contents": target_content} 
    
    else: 
        structure = focus_node['structure']

        control_center_contents = {k: v['user_instruction'] for k,v in control_center.items()}
        structure['initial_contents'] = "\n\n".join(["{}: {}".format(k,v) for k,v in control_center_contents.items()])
        mome.update_section(focus_node['path'], "Section structure", momeutils.j_deco(structure))

    # # PROPAGATION
    if is_root(config, focus_node['path']):
        momeutils.crline('Root node, no backprop')
        save_config(config, config_path)
        return

    hierarchy = collect_hierarchy_to_children(config, momeutils.bn(focus_node['path']), raw = True)
    _, _, name, _ = split_node_for_info(focus_node['path'])
    
    parent_node = os.path.join(config['interactive_graph_path'], hierarchy[-2] + ".md")
    parent_node_control = get_control_center(parent_node)
    parent_node_control[tmp_key_formatting(name, up_=True)]['user_instruction'] = structure['initial_contents']

    momeutils.crline('Propagating to parent node {}'.format(momeutils.bn(parent_node)))
    mome.update_section(parent_node, "Control center", momeutils.j_deco(parent_node_control))

    save_config(config, config_path)


def push_specific_parameter(config_path= None, **kwargs): 
    config= load_config(config_path, **kwargs)
    param_key = config['paragraph_parameter']
    param_content = config['paragraph_parameter_content']
    focus_node = get_focus_node(config)

    # collect all children leaf nodes 
    if is_leaf(focus_node['path']):
        target_children = [focus_node['path']]
    else: 
        target_children = collect_all_compilable_children(focus_node_path=focus_node['path'])

    psp_params = [
        (run_push_specific_parameter, (c, param_key, param_content)) for c in target_children
    ]
    psp_results = momeutils.mapper(psp_params)
    if max(psp_results) < 7: 
        momeutils.crline('Nodes are not relevant')
    else:
        max_val = max(psp_results)
        occurences = [i for i, v in enumerate(psp_results) if v == max_val]
        if len(occurences) > 1: 
            best_node_idx = occurences[1]# comparison_function([os.path.join(config['interactive_graph_path'], target_children[o] + ".md") for o in occurences])
        else: 
            best_node_idx = occurences[0]
        best_node = os.path.join(config['interactive_graph_path'], momeutils.bn(target_children[best_node_idx]) + ".md")
        control_center = get_control_center(best_node)
        control_center[param_key] = param_content
        mome.update_section(best_node, "Control center", momeutils.j_deco(control_center))

    save_config(config, config_path)        

def run_push_specific_parameter(node, param_key, param_content):

    control_center = get_control_center(node)
    lvl, part, name, _ = split_node_for_info(node) 
    momeutils.crline('Considering node: {}'.format(momeutils.bn(node)))
    result = evaluate_parameter_relevance("Section name: {}\nTarget key: {}\nContents: \n{}".format(name, param_key, json.dumps(control_center, indent = 4)), param_content)
    return result



    # for c in target_children: 
    #     control_center = get_control_center(c)
    #     lvl, part, name, _ = split_node_for_info(c) 
    #     momeutils.crline('Considering node: {}'.format(c))
    #     result = evaluate_parameter_relevance("Section name: {}\nTarget key: {}\nContents: \n{}".format(name, param_key, json.dumps(control_center, indent = 4)), param_content)


        # p = "In the context of automated long-form text generation, we are looking to enhance the current leaf paragraph with a {}. The initial (fixed) context is described by the following parameters: \n{}\n\nFrom a high-level standpoint, overlooking multiple paragraph, the user has requested the following content to be included \n{}\n\nBased on the context, is the user provided input relevant or should it be kept for other places ? Answer with a short rationale and then conclude with an integer score from 0 to 10, (10 being perfect fit).".format(
        #     param_key, json.dumps(control_center, indent = 4), "{} --> {}".format(param_key, param_content))
        # # input(p)
        # print('Considering node: {}'.format(c))
        # result = momeutils.basic_task(p, model = 'g4o')

def check_matching_key(control_center, control_key):

    matching = [k for k in control_center.keys() if k.lower().strip().startswith(control_key.lower().strip())]
    if len(matching) == 0: 
        raise ValueError(f"Control key {control_key} not found in control contents")
    return matching[0]

def run_review(config_path = None, **kwargs):
    config = load_config(config_path, **kwargs)
    focus_node = get_focus_node(config)
    review_contents = config['review_contents']
    control_center = focus_node['control']

    control_key = check_matching_key(control_center, config['control_key'])


    if is_leaf(focus_node['path']): 
        momeutils.crline('Is leaf node ! Process undefined')
        return 
    else: 
        current_contents = control_center[control_key]['user_instruction']
        out = adjust_text_based_on_review(current_contents, review_contents)
        control_center[control_key]['user_instruction'] = out
        mome.update_section(focus_node['path'], "Control center", momeutils.j_deco(control_center))
        save_config(config, config_path)


def refill(config_path = None, **kwargs): 

    """
    Refill is supposed to be used after a few sections have been filled and there is new 
    and additional information in the initial_contents of a section structure. 
    In this case, we might want to leverage those to enhance downstream sections.
    """

    config = load_config(config_path, **kwargs)
    focus_node = get_focus_node(config)
    
    if is_leaf(focus_node['path']):
        momeutils.crline('Leaf node, nothing to do')
        return
    
    control_center = focus_node['control']  
    key_id = check_matching_key(control_center, config['control_key'])
    section_structure = focus_node['structure']
    current_contents = control_center[key_id]['user_instruction']
# def enhance_report_section(section_identifier = None, current_contents = None, initial_global_contents = None, subsection_titles = None): 
    results = enhance_report_section(section_identifier = key_id, 
                                    current_contents = current_contents,
                                    initial_global_contents = section_structure['initial_contents'],
                                    subsection_titles = section_structure['subs_titles'])
    
    # momeutils.unconstrained_task("In the process of writing a report, new information has been obtained and the contents of a given section should now be enhanced. In particular, we are focusing on section: {}\n Current contents for the section {}\n{}. As mentionned, user has enhanced the initial global contents with the following information: \n\n{}\n\nBased on this new information, rewrite/improve the current overview contents of the section to include relevant pieces of the new information (there are some other sections ({}) as well, so stick to the relevant aspects to avoid redundancy).".format(key_id, key_id, current_contents, section_structure['initial_contents'], " ,".join(section_structure['subs_titles'])), model = "g4o")
    # input("ok  ? ")
    control_center[key_id]['user_instruction'] = results
    mome.update_section(focus_node['path'], "Control center", momeutils.j_deco(control_center))
    save_config(config, config_path)


def enhance_control_key(config_path=None, **kwargs):
    config = load_config(config_path, **kwargs)
    control_key = config['control_key']
    control_contents = config['control_contents']
    
    focus_node = get_focus_node(config)
    
    control_center = focus_node['control']

    if is_leaf(focus_node['path']):
        key_id= 'content'
        # TMP ! SHOULD HAVE SPECIALIZED AGENT
        out = enhance_section_content("",
                                      "", 
                                      split_node_for_info(focus_node['path'])[2],
                                      control_center[key_id],
                                     control_contents 
        )
        control_center[key_id] = out
    else: 
        section_structure = focus_node['structure'] 
        # Find the key that starts with the control_key
        key_id = check_matching_key(control_center, control_key)

        # p = "Initial complete content before split: {}\n\nTarget section: {}\n\nInitial content assigned to section: {}\n\nUser request: {}\n\nEnhance the Initial content assigned to section taking into account the global context and with particular attention to the user request.".format(section_structure['initial_contents'],
        #                                                  control_key, control_center[key_id], control_contents)
        # # out = momeutils.basic_task(p, model = "g4o")
        # def enhance_section_content(high_level_text = None, section_names = None, target_section = None, initial_section_content = None, user_instruction = None): 
        out = enhance_section_content(section_structure['initial_contents'], 
                                    "\n".join([""] + section_structure['subs_titles']), 
                                    key_id,
                                    control_center[key_id]['user_instruction'], 
                                    control_contents)
        control_center[key_id]['user_instruction'] = out
    mome.update_section(focus_node['path'], "Control center", momeutils.j_deco(control_center))
    save_config(config, config_path)   



def more_contents(config_path = None, **kwargs):
    config = load_config(config_path, **kwargs)

    focus_node_path = os.path.join(config['interactive_graph_path'], config['focus_node'] + ".md")
    if config['user_instruction'] is not None and config['user_instruction'].strip() != '': 
        hierarchy = collect_hierarchy_to_focus_node(config)
        section_structure =get_section_structure(focus_node_path) 
        subs = get_section_structure(focus_node_path)['subs_titles']
        p = "As an efficient AI writing assistant, you are to provide insightful complementary ideas to the user based on the following information. \n Section location in the document hierarchy : \n{}\n Focus section is {}\n Subsections in this section are: \n{}\n\n The user has provided the following initial text: \n{}\n\nFinally, the user has also provided the following instruction: {}\n\nBased on this global view, rewrite/expand/enhance the initial provided text taking into account the additional instruction.".format("\n".join(["\t" + h for h in hierarchy]), config['focus_node'].split('_')[-2], "\n".join(["\t* " + s for s in subs]), section_structure['initial_contents'], config['user_instruction'])
        out = momeutils.basic_task(p, model = 'g4o')
        # valid = momeutils.u_valid()
        valid = True
        if valid: 
            section_structure = update_existing_structure(section_structure, initial_contents = out)  
            mome.update_section(focus_node_path, "Section structure", momeutils.j_deco(section_structure))
    save_config(config, config_path)  
            

def is_leaf(current_node, target_tag = "results"): 
    tags= mome.get_file_tags(current_node)
    return target_tag in tags

def collect_all_compilable_children(focus_node_path, target_tag = ["results"]):
    node_dynasty = mome.collect_dynasty_paths(focus_node_path, include_root=False, preserve_hierarchy=False)
    _, result_nodes = mome.collect_node_contents(os.path.dirname(focus_node_path), tags = target_tag, return_paths = True)
    to_compile = [n for n in node_dynasty if n in result_nodes]
    return to_compile

def compile(config_path=None, **kwargs):
    config = load_config(config_path, **kwargs)
    focus_node_path = os.path.join(config['interactive_graph_path'], config['focus_node'] + ".md")
    if is_leaf(focus_node_path): 
        to_compile = [focus_node_path]
    else: 
        to_compile = collect_all_compilable_children(focus_node_path)

    # momeutils.dj(to_compile)
    # RUNNING THE ACTUAL COMPILATION
    # SEQUENTIAL COMPILATION
    # compilation_results = {}
    # for c in to_compile:
    #     lvl, part, name, hash_ = split_node_for_info(c)
    #     current_result = compile_node(c, strategy = "neighbors", config = config)
    #     compilation_results[name] = current_result
    #     compilation_results[name]['lvl'] = lvl
    #     compilation_results[name]['part'] = part
    #     compilation_results[name]['name'] = name
    #     compilation_results[name]['hash'] = hash_

    # momeutils.dj(compilation_results)



    # print('You forgot to uncomment'.upper()*20)
    # compilation_results = json.load(open('tmp.json'))
    compilation_results = {}

    # PARALLELIZING THE COMPILATION
    func_arg_list = [
        (compile_node, (c, "Control center", "neighbors", config)) for c in to_compile
    ]
    parallel_results = momeutils.mapper(func_arg_list)
    
    # AGREGATING RESULTS AGAIN 
    for i,c in enumerate(to_compile):   
        current_result = parallel_results[i]
        lvl, part, name, hash_ = split_node_for_info(c)
        
        compilation_results[name] = current_result

        compilation_results[name]['lvl'] = lvl
        compilation_results[name]['part'] = part
        compilation_results[name]['name'] = name
        compilation_results[name]['hash'] = hash_

    # with open('tmp.json', 'w') as f:    
    #     json.dump(compilation_results, f, indent = 4)

    update_compilation_results(config, compilation_results)
    # input(' ok ')
    produce_latex(config)
    produce_word(config)
    fill_obsidian(config, compilation_results)

def produce_latex(config):
    if not os.path.exists(config['report_path']):
        os.makedirs(config['report_path'])
    contents = json.load(open(config['results_path']))['results']
    latex_content = []

    def traverse_structure(structure, level):
        if isinstance(structure, dict):
            for key, value in structure.items():
                if level == 0:
                    latex_content.append(f"\\section{{{tmp_key_formatting(key, up_ = True)}}}")
                elif level == 1:
                    latex_content.append(f"\\subsection{{{tmp_key_formatting(key, up_ = True)}}}")
                elif level == 2:
                    latex_content.append(f"\\subsubsection{{{tmp_key_formatting(key, up_ = True)}}}")
                else:
                    latex_content.append(f"\\paragraph{{{key}}}")
                
                traverse_structure(value, level + 1)
        elif isinstance(structure, list):
            for item in structure:
                traverse_structure(item, level)
        elif isinstance(structure, str):
            latex_content.append(structure)

    traverse_structure(contents, 0)
    latex_result = open(config['report_details']['template']).read()
    latex_result = latex_result.replace('MY_TITLE', config['report_details']['title']).replace('MY_AUTHOR', config['report_details']['author']).replace('MY_CONTENTS', "\n\n".join(latex_content))
    with open(os.path.join(config['report_path'], 'report.tex'), 'w') as f:
        f.write(latex_result)

    # RUN COMPILATION 
    subprocess.run(['pdflatex','-interaction=nonstopmode' ,'-output-directory', config['report_path'], os.path.join(config['report_path'], 'report.tex')])



def produce_word(config):
    if not os.path.exists(config['report_path']):
        os.makedirs(config['report_path'])
    
    contents = json.load(open(config['results_path']))['results']
    document = Document()

    def traverse_structure(structure, level):
        if isinstance(structure, dict):
            for key, value in structure.items():
                if level == 0:
                    document.add_heading(tmp_key_formatting(key, up_=True), level=1)
                elif level == 1:
                    document.add_heading(tmp_key_formatting(key, up_=True), level=2)
                elif level == 2:
                    document.add_heading(tmp_key_formatting(key, up_=True), level=3)
                elif level == 3:
                    document.add_heading(tmp_key_formatting(key, up_=True), level=4)
                elif level == 4:
                    document.add_heading(tmp_key_formatting(key, up_=True), level=5)
                else:
                    document.add_paragraph(key, style='BodyText')
                
                traverse_structure(value, level + 1)
        elif isinstance(structure, list):
            for item in structure:
                traverse_structure(item, level)
        elif isinstance(structure, str):
            document.add_paragraph(structure)


    # Add title and author
    document.add_heading(config['report_details']['title'], level=0)
    document.add_paragraph(f"Author: {config['report_details']['author']}")

    traverse_structure(contents, 0)


    # Save the document
    document_path = os.path.join(config['report_path'], 'report.docx')
    document.save(document_path)



def find_node_in_structure(structure, lvl, part, name, current_lvl=0, current_part=0, path=[]):
        # CAREFUL ! THERE IS A VULNERABILITY IN FIND_NODE --> NODES THAT DIFFER ONLY BY THE HASH CAN BE MIXED UP 

        # print('Current path: {} - Looking for lvl {} and part {} - Current lvl {} and part {}'.format(path, lvl, part, current_lvl, current_part))

        # Check if the current level and part match the target
        if current_lvl == lvl and current_part == part:
            # input('here')
            if isinstance(structure, str) and name.strip() == structure.strip():
                # print('Found node {} at path {} - Structure: {}'.format(name, path, structure))
                # input(' ok ? ? ? ? ')
                return path    
            else: 
                if isinstance(structure, dict): 
                    if structure[list(structure.keys())[0]][0] == "To Fill" and path[-1] == name: 
                        return path
                elif isinstance(structure, list):
                    # momeutils.dj({"Structure": structure, "Name": name, "Path": path})
                    if structure[0] == "To Fill" and path[-1] == name: 
                        return path 
                    elif path[-1] == name: 
                        return path
            
        # Recursively search through dictionaries
        if isinstance(structure, dict):
            # print('Is dict')
            for key, value in structure.items():
                # print('Checking dict item {}'.format(key))
                result = find_node_in_structure(value, lvl, part, name, current_lvl + 1, current_part, path + [key])
                if result:
                    return result
        # Recursively search through lists
        elif isinstance(structure, list):
            # print('Is list')
            for index, item in enumerate(structure):
                # print('Checking list item {}'.format(index))
                result = find_node_in_structure(item, lvl, part, name, current_lvl, index, path + [index])
                if result:
                    return result
        return None

def update_compilation_results(config, compilation_results):

    # Initialize the results file if it doesn't exist
    if not os.path.exists(config['results_path']):
        os.makedirs(os.path.dirname(config['results_path']), exist_ok=True)
        with open(config['results_path'], 'w') as f:
            json.dump({"structure": config['report_structure'], "results": config['report_structure']}, f, indent=4)
    
    # Load existing results
    with open(config['results_path'], 'r') as f:
        existing_results = json.load(f)
    
    # Update the results
    for k, result in compilation_results.items():
        lvl = result['lvl']
        part = result['part']
        name = result['name']
        target_hash = result['hash']
        
        print('\n\nLooking for node with lvl {} and part {} - Name {}'.format(lvl, part, name))
        print('='*10)
        
        # Find the corresponding node in the structure
        print('='*10)
        formatted_name = tmp_key_formatting(name, up_ = True).strip()

        hierarchy = find_node_in_structure(config['report_structure'], lvl, part, formatted_name, current_lvl=0, current_part=0)
        # if name.startswith('Significant'):
        #     input("above")
        if hierarchy:
            # Navigate to the correct place in the results
            current_level = existing_results['results']
            for h in hierarchy[:-1]:
                current_level = current_level[h]

            # Update the results based on the structure type
            if isinstance(current_level, dict):
                if isinstance(current_level[hierarchy[-1]], list):
                    current_level[hierarchy[-1]][part] = {name: result['result']}
                elif isinstance(current_level[hierarchy[-1]], str):
                    current_level[hierarchy[-1]] = {name: result['result']}
            elif isinstance(current_level, list):
                current_level[part] = {name: result['result']}
    
    # Save the updated results
    with open(config['results_path'], 'w') as f:
        json.dump(existing_results, f, indent=4)

    
def fill_obsidian(config, compilation_results):

    focus_node_path = os.path.join(config['interactive_graph_path'], config['focus_node'] + ".md")
    # FORMATTING 
    if not is_leaf(focus_node_path): 
        compiled_text = []
        current_lvl = compilation_results[list(compilation_results.keys())[0]]['lvl']
        for k in compilation_results.keys(): 

            # TAKING INTO ACCOUNT SECTION CHANGES WHEN COMPILING
            if compilation_results[k]['lvl'] > current_lvl: 
                # get the new section: its lvl-th sub_titles from the section structure
                new_section = get_section_structure(focus_node_path)['subs_titles'][current_lvl]
                compiled_text.append(fill_section_template({"name": new_section, "lvl": current_lvl}, only_section = True))
                current_lvl = compilation_results[k]['lvl']

            compiled_text.append(fill_section_template(compilation_results[k]))
        mome.update_section(focus_node_path, "Results", "\n\n".join(compiled_text))
        # mome.update_section(focus_node_path, "Results", "\n\n".join([compilation_results[k]['result'] for k in compilation_results.keys()]))
    momeutils.crline('Compilation results: \n{}'.format(json.dumps([compilation_results[k]['compiled'] for k in compilation_results.keys()], indent = 4)))

def fill_section_template(current_result, only_section = False):
    # ASSUMING .section_template.txt exists in the folder 
    
    template = open(os.path.join(os.path.dirname(__file__), '.section_template.txt'), 'r').read()
    section_type = "section" if current_result['lvl'] == 0 else "sub" * current_result['lvl'] + "section"
    if only_section:
        template = template.strip().split('\n')[0]
        template = template.replace('ZENAME', current_result['name']).replace('section_type', section_type)
        return template
    return template.replace('section_type', section_type).replace('ZENAME', current_result['name']).replace('ZEPART', str(current_result['part'])).replace('ZERESULT', current_result['result'])


def compile_node(c, target_section = "Control center", strategy = "base", config = None): 
    
    if strategy == "base":
        paragraph_controls = momeutils.parse_json(mome.get_node_section(c, target_section))
        if paragraph_controls['content'] is None: 
            return {'compiled': False, 'result': 'No content to compile for {}'.format(momeutils.bn(c).split('_')[-2])}
        
        paragraph_reprez = {k:v for (k, v) in paragraph_controls.items() if (k != "nb_paragraphs" and v is not None)}
        paragraph_reprez = json.dumps(paragraph_reprez, indent = 4)
        test = generate_paragraph_subcontents(paragraph_reprez, paragraph_controls['nb_paragraphs'])

        result = paragraph_writer2(test)
    elif strategy == "neighbors": 

        lvl, part, name, _ = split_node_for_info(c)
        target_control = momeutils.parse_json(mome.get_node_section(c, target_section))
        neighbors = find_neighbors(config, momeutils.bn(c))
        neighborhood = []
        for i in range(len(neighbors)+1): 
            if i == part: 
                contents = {"Current paragraph": name, 
                            "Position": i,
                            "Controls": momeutils.parse_json(mome.get_node_section(c, target_section))}
            else: 
                target_node = [n for n in neighbors if n[0] == i][0][1]
                
                n_lvl, n_part, n_name, _ = split_node_for_info(target_node)
                n_path = os.path.join(config['interactive_graph_path'], target_node + ".md")
                if is_leaf(n_path):
                    controls = {"content": momeutils.parse_json(mome.get_node_section(n_path, target_section))}
                    control_contents = controls['content']
                else: 
                    momeutils.crline('HERE {}\n'.format(n_path) * 10)
                    controls = momeutils.parse_json(mome.get_node_section(n_path, target_section))
                    # momeutils.dj(controls)
                    # control_contents = "\n* ".join([["The neighbor section {} contains several subsections".format(tmp_key_formatting(n_name, up_ = True))] + ["{}: {}".format(tmp_key_formatting(k, up_ =True), v['user_instruction']) for k,v in controls.items()]])
                    control_contents = "\n* ".join(
                        ["The neighbor section {} contains several subsections".format(tmp_key_formatting(n_name, up_=True))] +
                        ["{}: {}".format(tmp_key_formatting(k, up_=True), v['user_instruction']) for k, v in controls.items()]
                    )
                    # momeutils.crline(control_contents)
                    # control_contents = controls[tmp_key_formatting(n_name, up_ = True)]['user_instruction']
                    # print('OKAY for {}\n'.format(n_name)*10)
                contents = {"Neighbor section": n_name, 
                            "Position": "Current section + {}".format(n_part),
                            "Controls": control_contents}
                
            neighborhood.append(contents)
        neighborhood.append('Again, to reiterate, the current paragraph to focus on is {} at position {}. Your answer must only develop this part, the rest is provided as context to avoid reintroduce already existing concepts and ensure smooth transitions between topics.'.format(tmp_key_formatting(name, up_ = True), part))
        control_archi = generate_paragraph_subcontents(neighborhood, target_control['nb_paragraphs'])
        additional_guidance = "\n\nAdditional guidance: \n {}".format("\n * ".join([""] + ["{}: {}".format(k.upper(), target_control[k]) for k in ['specific_viewpoint', "examples", "visuals", "references"] if target_control[k] is not None])) if (target_control['examples'] is not None or target_control['specific_viewpoint'] is not None or target_control['references'] is not None) else ""
        # momeutils.crline(additional_guidance)
        control_archi_enhanced = "\n".join(["Structure", "\n".join(control_archi), additional_guidance])

        result = paragraph_writer2([control_archi_enhanced])

        # momeutils.dj({"Neigh": neighborhood, 
        # "archi": control_archi, 
        # "result": result})

    mome.update_section(c, "Results", result)
    return {'compiled': True, 'result': result}

def get_focus_node(config, targets = ['Section structure', 'Control center']): 

    focus_node_path = os.path.join(config['interactive_graph_path'], config['focus_node'] + ".md")
    result = {"path": focus_node_path}
    for t in targets:
        if mome.check_section(focus_node_path, t): 
            result[t.lower().replace('section', '').replace('center', '').strip()] = momeutils.parse_json(mome.get_node_section(focus_node_path, t))
    
    return result

def tmp_key_formatting(s, up_ = False):
    result = []
    for char in s:
        if char.isupper():
            if result:  # Avoid adding a space at the beginning
                result.append(' ')
            if up_:
                result.append(char)
            else: 
                result.append(char.lower())
        else:
            result.append(char)
    return ''.join(result)


def find_next_level_structure(report_structure, hierarchy):

    next_level_structure = report_structure

    for level in hierarchy[1:]: 
        # momeutils.crline("Target level - Level {} - Formatted: {}".format(level, tmp_key_formatting(level)))
        if isinstance(next_level_structure, dict):
            # momeutils.crline('Looking for dict {}'.format(tmp_key_formatting(level)))   
            for k in next_level_structure.keys():
                # print('Checking key {}'.format(k))
                if k.lower() == tmp_key_formatting(level):
                    next_level_structure = next_level_structure[k]
                    # momeutils.crline('Found dict {} - Next structure {}\n Structure type :{} \n\n'.format(k, json.dumps(next_level_structure, indent = 4), type(next_level_structure)))
                    # momeutils.crline('Structure:\n{}'.format(next_level_structure))
                    break
        elif isinstance(next_level_structure, list):
            # momeutils.crline('Looking for list {}'.format(tmp_key_formatting(level)))
            for i, v in enumerate(next_level_structure):
                # print('Checking list id : {} - Value : {}'.format(i, v))
                str_v = v if isinstance(v, str) else list(v.keys())[0]
                if str_v.lower() == tmp_key_formatting(level):
                    # momeutils.dj(next_level_structure)
                    next_level_structure = next_level_structure[i]
                    if isinstance(next_level_structure, dict):
                        next_level_structure = list(next_level_structure.values())[0]
                    # momeutils.crline('Found list {} - Next structure {}\n Structure type :{} \n\n'.format(str_v, json.dumps(next_level_structure, indent = 4), type(next_level_structure)))
                    # momeutils.crline('Structure:\n{}'.format(next_level_structure))
                    break                   
    
    # momeutils.dj(next_level_structure)
    children = []
    if isinstance(next_level_structure, dict):
        next_level_structure = list(next_level_structure.values())[0]
    for c in next_level_structure: 
        if isinstance(c, dict): 
            children.append(list(c.keys())[0])
        elif isinstance(c, str): 
            children.append(c)
        else: 
            raise TypeError('Unrecognized type {} for next level structure {}'.format(type(c), c))  
    return children

def subs_titles_syntax_check(focus_node_contents):
    """
    Currently, only ensures that the subs_titles are in title case
    """
    focus_node_contents['structure']['subs_titles'] = [s.title() for s in focus_node_contents['structure']['subs_titles']]
    mome.update_section(focus_node_contents['path'], 'Section structure', momeutils.j_deco(focus_node_contents['structure']))


def is_root(config, focus_node_path):
    return config['current_hash'] == momeutils.bn(focus_node_path)

def structure_propagation(config_path = None, **kwargs): 
    config= load_config(config_path, **kwargs)
    focus_node = get_focus_node(config)
    hierarchy = collect_hierarchy_to_focus_node(config)

    # momeutils.dj(hierarchy)
    # momeutils.dj(focus_node['control'])
    
    subs_titles_syntax_check(focus_node)  

    # collect the initial contents (first level) from the report structure 
    if is_root(config, focus_node['path']):
        next_level_structure = list(config['report_structure'].keys())
    else: 
        next_level_structure = find_next_level_structure(config['report_structure'], hierarchy)
    # momeutils.dj(config['report_structure'])
    # momeutils.dj(next_level_structure)

    if next_level_structure != focus_node['structure']['subs_titles']: # if the user has updated something 
         
        # Based on the focus_node['structure']['subs_titles'], figure out how the downstream names must change for the already existing nodes (all existing nodes after those potentially added by the user will undergo a name change)
        to_change = [] 
        to_add= []
        for i in range(len(focus_node['structure']['subs_titles'])):
            if focus_node['structure']['subs_titles'][i] in next_level_structure:
                if next_level_structure.index(focus_node['structure']['subs_titles'][i]) != i:
                    to_change.append({'node' : focus_node['structure']['subs_titles'][i],
                                      'initial_index': next_level_structure.index(focus_node['structure']['subs_titles'][i]), 
                                      'current_index': i}) 
            else: 
                to_add.append({'node': focus_node['structure']['subs_titles'][i], 
                               'current_index': i})
        to_remove = [{"node": n, "current_index" : i} for i, n in enumerate(next_level_structure) if n not in focus_node['structure']['subs_titles']]        

        current_dynasty = format_dynasty((mome.collect_dynasty_paths(focus_node['path'], include_root=True, preserve_hierarchy=True)), keep_path= True)

        # momeutils.dj({'to_change': to_change, 
        #               'to_add': to_add, 
        #               "to_remove": to_remove, 
        #               "focus_node": focus_node['structure']['subs_titles'],
        #               "focus_path": focus_node['path']})
        
        current_control_center = focus_node['control']
        # momeutils.dj(current_control_center)

        # UPDATE EXISTING NODES  
        for node in to_change: 

            current_node_path = collect_path_from_formatted_dynasty(node['node'], current_dynasty['children'])
            lvl, part, name, hash = split_node_for_info(current_node_path)
            new_name = f"lvl{lvl}_part{node['current_index']}_{name}_{hash}"
            # ideally, we should also change children names but that would imply first changing leaves and then bringing information back up to ensure links stay valid
            # mome.obsidian_rename_node(current_node_path, new_name)   
            hierarchical_rename(current_node_path, new_name)         

            current_dynasty['children'][node['initial_index']]['path'] = os.path.join(os.path.dirname(current_node_path), f"{new_name}.md")
            
            # update links in the focus node 
            mome.replace_link(focus_node['path'], momeutils.bn(current_node_path), new_name)

        # Prepare missing nodes creation --> nodes should be added in 'pour' so that changes in the template can be handled 
        for node in to_add:
            current_contents = get_default_section_dict()
            # current_control_center= get_control_center(focus_node['path'])
            current_control_center = momeutils.insert_key_at_idx(current_control_center, 
                                                                 node['current_index'], 
                                                                 node['node'], 
                                                                 _value= current_contents)
            if momeutils.bn(focus_node['path']) == config['current_hash']:
                add_lvl = 0
                current_hash = config['current_hash']
                current_node_name = f"lvl{add_lvl}_part{node['current_index']}_{node['node'].title().replace(' ', '')}_{current_hash}.md"
                # current_dynasty['children'].insert(node['current_index'],
                #                                    {"path": 
                #                                     }
            else: 
                add_lvl = 1 + int(os.path.basename(focus_node['path']).split('_')[0].replace('lvl',''))
                current_hash = mome.get_short_hash(os.path.basename(focus_node['path']).split('.')[0], 15)
                current_node_name = f"lvl{add_lvl}_part{node['current_index']}_{node['node'].title().replace(' ', '')}_{current_hash}.md"
            current_dynasty['children'].insert(node['current_index'], 
                                                {'path': os.path.join(os.path.dirname(focus_node['path']), current_node_name), 
                                                    'children': []})
            # momeutils.dj(current_control_center)
            mome.update_section(focus_node['path'], 'Control center', momeutils.j_deco(current_control_center))

            # Updating links --> Adding the link too early introduces problems when parsing the hierarchy
            # mome.add_link_at_position(focus_node['path'], current_node_name , node['current_index'])

        # momeutils.dj(to_remove)
        # momeutils.dj(current_dynasty['children'])
        # Remove nodes
        for node in to_remove:
            # print(node['node'], list(current_control_center.keys()), node['node'] in list(current_control_center.keys()))
            node_path = collect_path_from_formatted_dynasty(node['node'], current_dynasty['children'])

            clean_dynasty(node_path, include_root=True)
            # also, to remove from control center 
            # current_control_center= focus_node['control']
            current_control_center.pop(node['node'])

            mome.update_section(focus_node['path'], 'Control center', momeutils.j_deco(current_control_center))
            # current_dynasty['children'].pop(node['current_index'])

            # Actually removing from the dynasty requires some gymnastics because the initial indices may be affected by previous operations
            node_to_remove = [n for n in current_dynasty['children'] if n['path'] == node_path][0]
            current_dynasty['children'].remove(node_to_remove)

            # And finally, pop from links 
            mome.remove_links(focus_node['path'], [momeutils.bn(node_path)]) 

        current_dynasty_copy = copy.deepcopy(current_dynasty)

        # SAVING THE FINAL REPORT STRUCTURE 
        config['report_structure'] = update_report_structure(config, current_dynasty_copy, focus_node['path'])
        save_config(config, config_path)

def hierarchical_rename(current_node_path, new_name):
    # Collect the current dynasty paths
    current_dynasty = format_dynasty(
        mome.collect_dynasty_paths(current_node_path, include_root=True, preserve_hierarchy=True),
        keep_path=True
    )

    new_hash = mome.get_short_hash(new_name, 15)

    # Update the current node
    update_node_and_links(current_node_path, new_name, new_hash)
    # Recursively update children
    for child in current_dynasty['children']:
        child_path = child['path']
        lvl, part, name, hash = split_node_for_info(child_path)
        new_child_name = f"lvl{lvl}_part{part}_{name}_{new_hash}"
        hierarchical_rename(child_path, new_child_name)

def update_node_and_links(node_path, new_name, new_hash):
    # Get current links
    current_links = mome.get_node_links(node_path)
    new_links_names = []

    for link in current_links:
        lvl, part, name, hash = split_node_for_info(link)
        new_links_names.append(f"lvl{lvl}_part{part}_{name}_{new_hash}")

    # Update links and rename node
    mome.update_link_section(node_path, new_links_names)
    mome.obsidian_rename_node(node_path, new_name)


# def hierarchical_rename(current_node_path, new_name):

#     # Collect the current dynasty paths
#     current_dynasty = format_dynasty(
#         mome.collect_dynasty_paths(current_node_path, include_root=True, preserve_hierarchy=True),
#         keep_path=True
#     )
#     print(json.dumps(current_dynasty, indent = 4))
#     print(new_name, mome.get_short_hash(new_name, 15)) # this is the new name
#     new_hash = mome.get_short_hash(new_name, 15)

#     # =======================================
#     # EXAMPLE FOR THE CURRENT NODE, THIS HAS TO BE DONE RECURSIVELY
#     current_links = mome.get_node_links(current_node_path)
#     new_links_names = []
#     for link in current_links:
#         lvl, part, name, hash = split_node_for_info(link)
#         new_links_names.append(f"lvl{lvl}_part{part}_{name}_{new_hash}")

#     # renaming 
#     mome.update_link_section(current_node_path, new_links_names)
#     mome.obsidian_rename_node(current_node_path, os.path.join(os.path.dirname(current_node_path), new_name + ".md"))
        


def update_report_structure(config, current_dynasty, focus_node_path):
    report_structure = config['report_structure']
   
    current_dynasty_dict = dynasty_to_report_structure(current_dynasty)
    formatted_hierarchy = [tmp_key_formatting(h).title().replace('  ', ' ') for h in collect_hierarchy_to_focus_node(config)[1:]]

    # print(current_dynasty_dict)
    # momeutils.dj(formatted_hierarchy)
    if is_root(config, focus_node_path):
        report_structure = current_dynasty_dict
    else: 
        momeutils.update_nested_dict(report_structure, formatted_hierarchy, current_dynasty_dict)
    
    return report_structure
        
def dynasty_to_report_structure(current_dynasty):
    report_structure = []

    for node in current_dynasty['children']:
        node_name = tmp_key_formatting(split_node_for_info(node['path'])[-2]).title()
        
        if node['children']:
            # If the node has children, recursively convert them
            child_structure = dynasty_to_report_structure(node)
            report_structure.append({node_name: child_structure})
        else:
            # If the node has no children, it's a leaf node
            report_structure.append(node_name)

    return report_structure


def collect_path_from_formatted_dynasty(node_name, current_dynasty): 
    # print(json.dumps(current_dynasty, indent=4))
    # momeutils.uinput("{}".format(json.dumps(current_dynasty, indent =4)))
    for nn in current_dynasty: 
        # momeutils.crline(nn['name'])
        if 'name' in nn.keys():
            if nn['name'].lower().strip() == node_name.lower().strip(): 
                return nn['path']
        
def split_node_for_info(path):

    # p = os.path.basename(path).split('.')[0]
    p = momeutils.bn(path)
    if '_' in p: 
        c = p.split('_')
        lvl = int(c[0].replace('lvl', ''))
        part = int(c[1].replace('part', ''))
        name = c[2]
        hash = c[3]
        return lvl, part, name, hash 
    else: 
        return 0,0,"",p
    


def node_expansion_colab(config_path = None, **kwargs): 
    config = load_config(config_path, **kwargs)
    focus_node_path = os.path.join(config['interactive_graph_path'], config['focus_node'] + ".md")
    
    parent_hash = mome.get_short_hash(focus_node_path, 15)
    # GETS THE JSON WITH EACH PART (CAN BE SECTION OR SUBSECTION)
    focus_node_contents = momeutils.parse_json(mome.get_node_section(focus_node_path)) # Collects the first section
    focus_node_control = momeutils.parse_json(mome.get_node_section(focus_node_path, "Control center")) 
    
    # check if keys are the same 
    assert list(focus_node_contents.keys()) == list(focus_node_control.keys())

    result_node = mome.add_node_to_graph(config['interactive_graph_path'],
                                        contents = {"Results": ""}, 
                                        tags = ['result_node'],
                                        parent_path = focus_node_path, 
                                        name_override = "results_" + parent_hash,
                                        use_hash = False)
    
    current_tag = mome.get_file_tags(focus_node_path)[0] # CAREFUL FOR THE 0, IN CASE THERE ARE MULTIPLE TAGS AT SOME POINT 
    section_level = len(current_tag.split('_')) + 1
    added_nodes = []
    for i, k in enumerate(focus_node_contents.keys()): 
        # NAME CONVENTION IS: lvlI_partJ_name_parenthash
        n = add_subnode_from_contents_control(config, focus_node_path, focus_node_contents[k], focus_node_control[k], k, current_tag, section_level, i, parent_hash)
        added_nodes.append(n)
    
    config['last_nodes_added'] = added_nodes
    
    # enhance_root_structure(config)
    save_config(config, config_path)
    return config

def update_existing_structure(structure, **kwargs): 
    for k, v in kwargs.items(): 
        if k in structure.keys(): 
            structure[k] = v
    return structure

def initiate_control_center(**kwargs): 
    c = {
        "hierarchy_location": None, 
        "high_level_context": None, 
        "things_said_before": None, 
        "content": None,
        "where_this_is_going": None,
        "nb_paragraphs": 3,
        "examples": None, 
        "specific_viewpoint": None, 
        "references": None,
        "visuals": None 
    }

    for k, v in kwargs.items():
        if k in c.keys(): 
            c[k] = v
    return c 


def tmp_get_section_org(section_content):   
    sentences = section_content.split('.')

    ks = {" ".join(s.split()[-2:]).strip().replace(' ', '_'): s for s in sentences if len(s) > 0}
    return ks 


def add_subnode_from_contents_control(config, focus_node_path, contents, control, subname, current_tag, section_level, paragraph_id, parent_hash): 
    
    if control['template'].strip().lower() == "direct": # means that this is going to be directly used to write 
            tag = ['to_compile']
            focus_contents = momeutils.parse_json(mome.get_node_section(focus_node_path)) # Collects the first section
            t = initiate_control_center(content = contents, 
                                    things_said_before = "\n".join([focus_contents[kk] for kk in list(focus_contents.keys())[:list(focus_contents.keys()).index(subname)]]),
                                    where_this_is_going = "\n".join([kk for kk in list(focus_contents.keys())[list(focus_contents.keys()).index(subname)+1:]]),
                                    )
            node_contents = {"Structure": momeutils.j_deco(t), "Results" : ""}
    else: 
        section_org = tmp_get_section_org(contents)
        node_contents = {"Base contents": momeutils.j_deco(section_org), # AI RESULTS 
                            "Control center": momeutils.j_deco(setup_control_params(section_org)), 
                        "Section structure": momeutils.j_deco(setup_section_structure(section_org)),
                        }
        
        tag = ['sub_' + current_tag]

    new_part = mome.add_node_to_graph(config['interactive_graph_path'],
                                    contents = node_contents,
                                    tags = tag,
                                    parent_path = focus_node_path, 
                                    # node_prefix = k, 
                                    name_override = "lvl{}_part{}_{}_{}".format(section_level, paragraph_id, 
                                                                                re.sub(r'[^a-zA-Z0-9\s]', '', subname).strip().lower().replace(' ', '_'),
                                                                                parent_hash ),
                                    use_hash = False)
    return new_part

def do_section_struct_to_base_contents(focus_node_path): 
    structure_contents = momeutils.parse_json(mome.get_node_section(focus_node_path, "Section structure"))
    
    # TMP MANUAL FOR NOW !!! 
    computed_contents = {"_".join(s.strip().replace(',', '').split()[:2]).lower(): s for s in structure_contents['initial_contents'].split('.') if len(s.strip()) > 2}
    # ABOVE SHOULD BE A LLM CALL  

    # REMOVE NODE DYNASTY (CLEANING)
    clean_dynasty(focus_node_path, include_root = False)
    # dynasty = mome.collect_dynasty_paths(focus_node_path, include_root = False, preserve_hierarchy = False)
    

    mome.update_section(focus_node_path, "Base contents", momeutils.j_deco(computed_contents))
    # save_config(config, config_path)

    # HANDLES CONTROL CENTER 
    do_control_center_from_base_contents(focus_node_path)
    # control_center_from_base_contents(config_path)

def do_control_center_from_base_contents(focus_node_path):
    base_contents = momeutils.parse_json(mome.get_node_section(focus_node_path))    
    control_center = {k: get_default_section_dict() for k in base_contents.keys()}
    mome.update_section(focus_node_path, "Control center", momeutils.j_deco(control_center))

def get_default_section_dict(): 
    return {"user_instruction" : None, 
        "consistency_elements": None,   
        "additional_details": None, 
        "template": "default"}

def clean_dynasty(root, include_root, link_section = "Links"):
    """
    Removes children and cleans also the link sections, assuming include_root is false (otherwise, it is deleted) 
    """
    if root is None:
        momeutils.crline('Root node is None')
        return
    if not os.path.exists(root):
        momeutils.crline('Missing root node {}'.format(root))
        return
    dynasty = mome.collect_dynasty_paths(root, include_root = include_root, preserve_hierarchy = False)
    for node in dynasty: 
        os.remove(node)
    if not include_root: 
        mome.update_link_section(root, "", link_section = link_section)

def setup_control_center_from_hls(structure):
    subs_titles = gf_all(structure, show_structure=True)
    # momeutils.dj({"Subs titles": subs_titles})
    return setup_control_center(subs_titles)
 
def setup_control_center(named_templates): 

    base_format = {"user_instruction": None,
                   "consistency_elements": None, 
                   "additional_details": None,
                   "template": "default"}
    results = {}
    for i in range(len(named_templates)): 
        b = base_format.copy()
        if named_templates[i][1].strip() == "paragraph" or named_templates[i][1].strip() == "direct": 
            b['template'] = "direct"
        results[named_templates[i][0]] = b
    return results 

def setup_section_structure_from_hls(current_structure): 

    subs_titles = gf_all(current_structure, show_structure=True)
    initial_contents = ". ".join(["There is a {} about {}".format(s[1], s[0]) for s in subs_titles])
    return setup_section_structure(initial_contents = initial_contents, 
                                   subs_titles = [s[0] for s in subs_titles])

def setup_section_structure(**kwargs): 
    c = {"initial_contents": "", 
         "subs_titles": []}
    for k in kwargs.keys(): 
        if k in c.keys(): 
            c[k] = kwargs[k]
    return c

def setup_control_params(suggested_sections): 
    if isinstance(suggested_sections, list): 
        parts = list(suggested_sections)
    elif isinstance(suggested_sections, dict): 
        parts = list(suggested_sections.keys())
    else: 
        raise TypeError('Unrecognized type in setup control params for :{}'.format(type(suggested_sections)))
    
    return {part: get_default_section_dict() for part in parts}




if __name__ == "__main__": 

    config_path = os.path.join(os.path.dirname(__file__), "config_{}.json".format(os.path.basename(__file__).split('.')[0]))
    # a file containing debugging contents
    rationale_path = os.path.join(os.path.dirname(__file__), "contents_ainimals.json")
    rationales = json.load(open(rationale_path))
    make_graph(config_path)
    # updating root
    
    # config = json.load(open(config_path))

    # section_structure = get_section_structure(os.path.join(os.path.dirname(__file__), "sir_interactive_graph", config['current_hash'] + ".md"))
    # section_structure = update_existing_structure(section_structure, initial_contents = rationales['rationale'])
    
    # # UPDATING ROOT 
    # mome.update_section(os.path.join(os.path.dirname(__file__), "sir_interactive_graph", config['current_hash'] + ".md"), "Section structure", momeutils.j_deco(section_structure)) 

    # # UPDATING OPERATION TOPIC 
    # section_structure = get_section_structure(os.path.join(os.path.dirname(__file__), "sir_interactive_graph", "lvl0_part0_OperationTopic_2cf24dba5fb0a30" + ".md"))
    # # section_structure = update_existing_structure(section_structure, initial_contents = rationales['rationale'])
    # section_structure = update_existing_structure(section_structure, initial_contents = rationales['rationale'], subs_titles = rationales['different_subs_titles_in_operationtopic'])
    # mome.update_section(os.path.join(os.path.dirname(__file__), "sir_interactive_graph", "lvl0_part0_OperationTopic_2cf24dba5fb0a30" + ".md"), "Section structure", momeutils.j_deco(section_structure)) 

